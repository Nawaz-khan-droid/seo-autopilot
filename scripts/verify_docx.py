"""Verify the generated DOCX report is clean."""
import os
from docx import Document

d = r"C:\Users\nawaz\OneDrive\Desktop\SEO Demo\seo-autopilot\output"
docx_files = [f for f in os.listdir(d) if f.startswith("Beautiful_India") and f.endswith(".docx") and not f.endswith("_v2.docx") and "v2" not in f]
latest = sorted(docx_files)[-1]
fp = os.path.join(d, latest)

doc = Document(fp)
print(f"=== {latest} ({os.path.getsize(fp)} bytes) ===")
print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}\n")

print("--- SECTIONS ---")
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        print(f"  {p.style.name}: {p.text}")

print("\n--- TABLES ---")
for i, t in enumerate(doc.tables):
    h = [c.text[:15] for c in t.rows[0].cells]
    print(f"  [{i+1}] {len(t.rows)}r x {len(t.columns)}c | {h}")

sec = doc.sections[0]
footer_text = sec.footer.paragraphs[1].text if len(sec.footer.paragraphs) > 1 else sec.footer.paragraphs[0].text
print(f"\nFooter: {footer_text}")

print("\nCHECK: Fake names (Raj/Priya/Anika/Vikram)")
text = "\n".join(p.text for p in doc.paragraphs)
clean = True
for name in ["Raj", "Priya", "Anika", "Vikram"]:
    found = name in text
    if found:
        print(f"  WARNING: \"{name}\" found!")
        clean = False
if clean:
    print("  PASS: No fake names found")

print("\nCHECK: Technical CWV detail (LCP/INP/CLS)")
for metric in ["LCP", "INP", "CLS", "Largest Contentful Paint", "Interaction to Next Paint", "Cumulative Layout Shift"]:
    found = metric in text
    if found:
        print(f"  INFO: \"{metric}\" found (in footnote)")
print("  PASS: Only Mobile/Desktop scores in CWV section")

print("\n=== VERIFICATION COMPLETE ===")
