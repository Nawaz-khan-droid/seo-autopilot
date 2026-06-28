"""Monthly SEO Report — Client-facing DOCX (14-16 slide cards).

Balanced coverage: off-page (GMB, link building, PR), on-page (keyword
rankings, content), and technical SEO. Digital Marketing Agency → Client
voice. Every section includes explanatory bullet points so the client
understands what each data point means for their business.
"""

from __future__ import annotations

import io
import os
import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from report.facts import ReportFacts
from report.charts import make_distribution_hbar_chart, make_psi_column_chart

logger = logging.getLogger(__name__)

NAVY = RGBColor(0x0F, 0x27, 0x47)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
DARK = RGBColor(0x1E, 0x29, 0x3B)
BODY = RGBColor(0x33, 0x41, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
RED = RGBColor(0xDC, 0x26, 0x26)
TEAL = RGBColor(0x14, 0xB8, 0xA6)
GRAY = RGBColor(0x94, 0xA3, 0xB8)

FN = "Calibri"


def _shd(cell, c):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _cell(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FN
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align is not None: p.alignment = align

def _para(doc, text, bold=False, color=None, size=10, align=None, before=0, after=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FN
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p

def _bullets(doc, items, size=10):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(item)
        r.font.size = Pt(size)
        r.font.name = FN

def _ev(v):
    if v is None: return "\u2014"
    if hasattr(v, "is_available") and not v.is_available: return "\u2014"
    val = v.value if hasattr(v, "value") else v
    return str(val) if val is not None else "\u2014"

def _evi(v):
    try: return int(float(_ev(v)))
    except: return 0

def _img(doc, img_bytes, w=5.0):
    if not img_bytes or len(img_bytes) < 100: return
    s = io.BytesIO(img_bytes)
    doc.add_picture(s, width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def _cstatus(s):
    if s >= 90: return "Good", GREEN
    if s >= 50: return "Needs Work", AMBER
    return "Poor", RED

def _access_required(doc, service: str, detail: str = ""):
    """Insert an 'Access Required' notice box into the report."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ACCESS REQUIRED")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = AMBER
    run.font.name = FN
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Connect {service} to view data.")
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    r2.font.name = FN
    if detail:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(detail)
        r3.font.size = Pt(8)
        r3.font.color.rgb = GRAY
        r3.font.name = FN
    doc.add_paragraph()


def _sev_color(s):
    s = s.lower().strip()
    if s == "critical": return RED
    if s in ("high", "error"): return AMBER
    if s in ("medium", "warning"): return ACCENT
    return GRAY


def build_clients_report(facts: ReportFacts, output_path: str = "",
                          google_trends_data: dict | None = None) -> str:
    m = facts.metadata
    if not output_path:
        sn = (m.client_name or "Client").replace(" ", "_")
        sm = (m.report_month or "Report").replace(" ", "_")
        output_path = f"{sn}_{sm}_Report.docx"

    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = FN
    style.font.size = Pt(10)
    style.font.color.rgb = BODY
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.2

    for lvl, (sz, clr, sb) in {1: (16, NAVY, 4), 2: (13, ACCENT, 3), 3: (11, DARK, 2)}.items():
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = FN
        hs.font.size = Pt(sz)
        hs.font.color.rgb = clr
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(sb)
        hs.paragraph_format.space_after = Pt(2)

    ranks = facts.rankings
    total = len(ranks)
    total_display = facts.rankings_total_estimated or total
    improved = sum(1 for r in ranks if _evi(r.change) > 0)
    dropped = sum(1 for r in ranks if _evi(r.change) < 0)
    p1 = sum(1 for r in ranks if _evi(r.position) == 1)
    p2 = sum(1 for r in ranks if _evi(r.position) == 2)
    p3 = sum(1 for r in ranks if _evi(r.position) == 3)
    top3 = sum(1 for r in ranks if _evi(r.position) <= 3)
    top10 = sum(1 for r in ranks if _evi(r.position) <= 10)
    cm = _evi(facts.cwv.mobile_score)
    cd = _evi(facts.cwv.desktop_score)
    th = _evi(facts.technical.health_score)
    t = facts.technical

    gmb_total = len(facts.gmb_keywords)
    gmb_p1 = sum(1 for k in facts.gmb_keywords if k.positions and _evi(k.positions[-1]) == 1)
    gmb_top3 = sum(1 for k in facts.gmb_keywords if k.positions and _evi(k.positions[-1]) <= 3)

    # ════════════════════════════════════════════
    # CARD 1: COVER
    # ════════════════════════════════════════════
    for _ in range(5): doc.add_paragraph()
    if m.agency_name:
        _para(doc, m.agency_name.upper(), bold=True, color=ACCENT, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(doc, "", size=4)
    _para(doc, "MONTHLY SEO REPORT", bold=True, color=NAVY, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "\u2500" * 36, color=ACCENT, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "", size=4)
    if m.client_name:
        _para(doc, m.client_name, bold=True, color=NAVY, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    if m.report_month:
        _para(doc, m.report_month, color=BODY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "", size=8)
    gen = m.generated_at or datetime.now().strftime("%B %d, %Y")
    _para(doc, f"Generated: {gen}", color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "CONFIDENTIAL", bold=True, color=GRAY, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ════════════════════════════════════════════
    # EXECUTIVE NARRATIVE (LLM-generated, optional)
    # ════════════════════════════════════════════
    if facts.executive_narrative:
        doc.add_heading("Executive Summary", level=1)
        _para(doc, facts.executive_narrative, size=10, color=BODY)
        _para(doc, "", size=4)

    # ════════════════════════════════════════════
    # CARD 2: EXECUTIVE DASHBOARD
    # ════════════════════════════════════════════
    doc.add_heading("Executive Dashboard", level=1)
    _para(doc, "At-a-glance view of your search performance this period.", size=9, color=GRAY)

    kpis = [("Keywords Tracked", str(total_display) if total_display else "\u2014", NAVY, "0F2747"),
            ("#1 Rankings", str(p1) if total else "\u2014", GREEN if total else GRAY, "16A34A"),
            ("Top 3", str(top3) if total else "\u2014", ACCENT if total else GRAY, "2563EB"),
            ("Top 10", str(top10) if total else "\u2014", TEAL if total else GRAY, "14B8A6")]
    t1 = doc.add_table(rows=2, cols=4)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (lb, vl, tc, bg) in enumerate(kpis):
        _cell(t1.rows[0].cells[ci], lb, bold=True, color=WHITE, size=8)
        _shd(t1.rows[0].cells[ci], bg)
        _cell(t1.rows[1].cells[ci], vl, bold=True, color=tc, size=15)
        _shd(t1.rows[1].cells[ci], "F8FAFC")

    if total > 0:
        t2 = doc.add_table(rows=1, cols=3)
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, (lb, tc, bg) in enumerate([(f"+{improved} Improved", GREEN, "16A34A"),
                                            (f"{dropped} Dropped", RED, "DC2626"),
                                            (f"{total - improved - dropped} Stable", GRAY, "6B7280")]):
            _cell(t2.rows[0].cells[ci], lb, bold=True, color=tc, size=10)
            _shd(t2.rows[0].cells[ci], "F8FAFC")

    sub_cards = []
    if cm > 0:
        sub_cards.append((f"Mobile CWV: {cm}/100", GREEN if cm >= 90 else (AMBER if cm >= 50 else RED)))
    if cd > 0:
        sub_cards.append((f"Desktop CWV: {cd}/100", GREEN if cd >= 90 else (AMBER if cd >= 50 else RED)))
    if th > 0:
        sub_cards.append((f"Tech Health: {th}/100", GREEN if th >= 80 else (AMBER if th >= 50 else RED)))
    if gmb_total > 0:
        sub_cards.append((f"GMB Keywords: {gmb_total} tracked", ACCENT))

    if sub_cards:
        t3 = doc.add_table(rows=1, cols=len(sub_cards))
        t3.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, (lb, tc) in enumerate(sub_cards):
            _cell(t3.rows[0].cells[ci], lb, bold=True, color=tc, size=9)
            _shd(t3.rows[0].cells[ci], "F8FAFC")

    _para(doc, "", size=4)
    dashboard_insights = []
    has_any_data = total > 0 or cm > 0 or th > 0
    if has_any_data:
        if total > 0:
            dashboard_insights.append(f"Keywords: {total} tracked | #1: {p1} | Top 3: {top3} | Top 10: {top10}")
        if cm > 0:
            dashboard_insights.append(f"Mobile CWV: {cm}/100" + (f" (below 90 threshold)" if cm < 80 else ""))
        if th > 0:
            dashboard_insights.append(f"Tech health: {th}/100")
    else:
        dashboard_insights.append("No performance data collected this period. Configure data sources to see metrics here.")
    _bullets(doc, dashboard_insights)
    doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 3: CORE SEO METRICS (title, meta, H1)
    # ════════════════════════════════════════════
    si = facts.site_info
    has_core_seo = bool(si.title_tag or si.meta_description or si.h1_count > 0)
    if has_core_seo:
        doc.add_heading("Core SEO Metrics", level=1)
        _para(doc, "How your page's foundational SEO elements are configured for search engines and users.", size=9, color=GRAY)

        core_items = []
        if si.title_tag:
            title_len = len(si.title_tag)
            title_status = "Good" if 30 <= title_len <= 60 else "Needs Improvement"
            core_items.append(("Title Tag", f"{si.title_tag[:60]}{'...' if len(si.title_tag)>60 else ''} ({title_len} chars)", title_status))
        if si.meta_description:
            desc_len = len(si.meta_description)
            desc_status = "Good" if 120 <= desc_len <= 160 else "Needs Improvement"
            core_items.append(("Meta Description", f"{si.meta_description[:80]}{'...' if len(si.meta_description)>80 else ''} ({desc_len} chars)", desc_status))
        h1_status = "Good" if si.h1_count == 1 else ("Error" if si.h1_count == 0 else "Warning")
        core_items.append(("H1 Headings", f"{si.h1_count} H1 tag(s) found", h1_status))
        if si.h1_texts:
            core_items.append(("H1 Text", "; ".join(si.h1_texts[:3]), ""))
        if si.word_count:
            wc_status = "Good" if si.word_count >= 300 else "Thin Content"
            core_items.append(("Word Count", f"{si.word_count:,} words", wc_status))

        ct = doc.add_table(rows=1 + len(core_items), cols=3)
        ct.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Element", "Value", "Status"]):
            _cell(ct.rows[0].cells[ci], h, bold=True, color=WHITE, size=9)
            _shd(ct.rows[0].cells[ci], "0F2747")
        for ri, (el, val, st) in enumerate(core_items):
            _cell(ct.rows[ri+1].cells[0], el, bold=True, size=9)
            _cell(ct.rows[ri+1].cells[1], val, size=8)
            st_color = GREEN if "Good" in st else (AMBER if "Warning" in st or "Needs" in st else (RED if "Error" in st else GRAY))
            _cell(ct.rows[ri+1].cells[2], st if st else "\u2014", bold=True, color=st_color, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2 == 0:
                for ci in range(3): _shd(ct.rows[ri+1].cells[ci], "F8FAFC")

        core_insights = []
        if si.h1_count == 0:
            core_insights.append("No H1 tag found — every page should have exactly one H1 describing its primary topic.")
        elif si.h1_count > 1:
            core_insights.append(f"Multiple H1 tags detected ({si.h1_count}) — consolidate to a single H1 for optimal SEO.")
        if si.meta_description and (len(si.meta_description) < 120 or len(si.meta_description) > 160):
            core_insights.append(f"Meta description is {len(si.meta_description)} characters (ideal: 120-160). Adjust to avoid truncation in SERPs.")
        if si.title_tag and len(si.title_tag) > 60:
            core_insights.append(f"Title tag exceeds 60 characters — consider shortening for full SERP display.")
        if si.word_count and si.word_count < 300:
            core_insights.append(f"Page content is thin ({si.word_count} words). Aim for at least 300 words of substantive content.")
        if core_insights:
            _para(doc, "", size=4)
            _para(doc, "Recommendations:", bold=True, size=10, color=NAVY)
            _bullets(doc, core_insights)

    # ════════════════════════════════════════════
    # CARD 4: BACKLINKS & AUTHORITY
    # ════════════════════════════════════════════
    bl = facts.backlinks
    has_backlink_data = _ev(bl.total_backlinks) != "\u2014" or _ev(bl.ref_domains) != "\u2014"
    has_internal_data = _ev(bl.onpage_total_links) != "\u2014"
    has_opr_rating = _ev(bl.domain_rating) != "\u2014"
    awaiting_bl = bl.status == "AWAITING_DATA"
    bl_insights: list[str] = []
    if has_backlink_data or awaiting_bl or has_internal_data or has_opr_rating:
        doc.add_heading("Backlinks & Authority", level=1)
        _para(doc, "Your website's external link profile and internal link distribution.", size=9, color=GRAY)

    if has_backlink_data:
        bl_items = []
        if _ev(bl.total_backlinks) != "\u2014":
            bl_items.append(("Total Backlinks", _ev(bl.total_backlinks)))
        if _ev(bl.ref_domains) != "\u2014":
            bl_items.append(("Referring Domains", _ev(bl.ref_domains)))
        if _ev(bl.dofollow_count) != "\u2014":
            bl_items.append(("Follow Links", _ev(bl.dofollow_count)))
        if _ev(bl.nofollow_count) != "\u2014":
            bl_items.append(("Nofollow Links", _ev(bl.nofollow_count)))

        if bl_items:
            bt = doc.add_table(rows=1 + len(bl_items), cols=2)
            bt.alignment = WD_TABLE_ALIGNMENT.CENTER
            _cell(bt.rows[0].cells[0], "Metric", bold=True, color=WHITE, size=9)
            _cell(bt.rows[0].cells[1], "Value", bold=True, color=WHITE, size=9)
            _shd(bt.rows[0].cells[0], "0F2747")
            _shd(bt.rows[0].cells[1], "0F2747")
            for ri, (lb, vl) in enumerate(bl_items):
                _cell(bt.rows[ri+1].cells[0], lb, size=9)
                _cell(bt.rows[ri+1].cells[1], vl, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                if ri % 2 == 0:
                    _shd(bt.rows[ri+1].cells[0], "F8FAFC")
                    _shd(bt.rows[ri+1].cells[1], "F8FAFC")

        _para(doc, "", size=4)
        bl_insights = []
        if _ev(bl.total_backlinks) != "\u2014":
            bl_count = int(_ev(bl.total_backlinks).replace(",", "").replace("+", ""))
            if bl_count > 0:
                bl_insights.append(f"Your site has {_ev(bl.total_backlinks)} backlinks from {_ev(bl.ref_domains) or 'multiple'} domains.")

    if has_opr_rating and not has_backlink_data:
        _para(doc, "", size=4)
        opr_table = doc.add_table(rows=1, cols=2)
        opr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _cell(opr_table.rows[0].cells[0], "Domain Authority (OpenPageRank)", bold=True, size=9)
        _cell(opr_table.rows[0].cells[1], _ev(bl.domain_rating), size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(opr_table.rows[0].cells[0], "F8FAFC")
        _shd(opr_table.rows[0].cells[1], "F8FAFC")
        _para(doc, "Upload a backlink CSV from Ahrefs, Semrush, or Moz for a full profile.", size=8, color=GRAY)

    if awaiting_bl:
        _para(doc, "", size=4)
        ct = doc.add_table(rows=1, cols=1)
        ct.alignment = WD_TABLE_ALIGNMENT.CENTER
        _shd(ct.rows[0].cells[0], "FFF8E1")
        _cell(ct.rows[0].cells[0], "", size=4)
        p = ct.rows[0].cells[0].paragraphs[-1]
        r = p.add_run("Action Required: Complete Your Backlink Profile")
        r.bold = True; r.font.size = Pt(10); r.font.name = FN
        r.font.color.rgb = RGBColor(0x8D, 0x6E, 0x00)
        p2 = ct.rows[0].cells[0].add_paragraph()
        r2 = p2.add_run(
            "Backlink data was not detected during this automated audit. "
            "To see your full backlink profile, export a CSV from your preferred "
            "tracking tool (Ahrefs, Semrush, Moz, or Ubersuggest) and upload it "
            "via the Data Uploads dashboard. Required columns: total_backlinks, "
            "ref_domains, dofollow, nofollow, domain_rating, source."
        )
        r2.font.size = Pt(9); r2.font.name = FN
        r2.font.color.rgb = RGBColor(0x8D, 0x6E, 0x00)

    if has_internal_data:
        _para(doc, "", size=4)
        il_total = _evi(bl.onpage_total_links)
        il_int = _evi(bl.onpage_internal_links)
        il_ext = _evi(bl.onpage_external_links)
        if il_total > 0:
            _para(doc, "Internal Link Distribution:", bold=True, size=10, color=NAVY)
            ilt = doc.add_table(rows=2, cols=3)
            ilt.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, (lb, vl) in enumerate([("Total", str(il_total)), ("Internal", str(il_int)), ("External", str(il_ext))]):
                _cell(ilt.rows[0].cells[ci], lb, bold=True, color=WHITE, size=9)
                _shd(ilt.rows[0].cells[ci], "0F2747")
                _cell(ilt.rows[1].cells[ci], vl, bold=True, color=ACCENT, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
                _shd(ilt.rows[1].cells[ci], "F8FAFC")
            if il_int > 0:
                bl_insights.append(f"{il_int} internal links help distribute page authority and guide users through your site architecture.")

    if awaiting_bl or has_backlink_data:
        if bl_insights:
            _para(doc, "", size=4)
            _bullets(doc, bl_insights)

    # ════════════════════════════════════════════
    # CARD 5: MARKETING & USABILITY
    # ════════════════════════════════════════════
    has_mktg = (si.has_og_tags.is_available or si.has_robots_txt.is_available or
                si.has_sitemap_xml.is_available)
    if has_mktg:
        doc.add_heading("Marketing & Usability", level=1)
        _para(doc, "Social sharing, crawl access, and technical foundation checks.", size=9, color=GRAY)

        mktg_items = []
        if si.has_og_tags.is_available:
            mktg_items.append(("Open Graph Tags", _ev(si.has_og_tags)))
        if si.has_robots_txt.is_available:
            mktg_items.append(("robots.txt", _ev(si.has_robots_txt)))
        if si.has_sitemap_xml.is_available:
            mktg_items.append(("XML Sitemap", _ev(si.has_sitemap_xml)))

        if mktg_items:
            mt = doc.add_table(rows=1 + len(mktg_items), cols=2)
            mt.alignment = WD_TABLE_ALIGNMENT.CENTER
            _cell(mt.rows[0].cells[0], "Check", bold=True, color=WHITE, size=9)
            _cell(mt.rows[0].cells[1], "Status", bold=True, color=WHITE, size=9)
            _shd(mt.rows[0].cells[0], "0F2747")
            _shd(mt.rows[0].cells[1], "0F2747")
            for ri, (lb, vl) in enumerate(mktg_items):
                _cell(mt.rows[ri+1].cells[0], lb, size=9)
                clr = GREEN if "Yes" in vl or "Found" in vl else (AMBER if "No" in vl else GRAY)
                _cell(mt.rows[ri+1].cells[1], vl, size=9, bold=True, color=clr, align=WD_ALIGN_PARAGRAPH.CENTER)
                if ri % 2 == 0:
                    _shd(mt.rows[ri+1].cells[0], "F8FAFC")
                    _shd(mt.rows[ri+1].cells[1], "F8FAFC")

        _para(doc, "", size=4)
        mktg_insights = []
        if _ev(si.has_og_tags) == "Yes":
            mktg_insights.append("Open Graph tags are properly configured — your pages will display rich previews on Facebook, LinkedIn, and Twitter/X.")
        elif _ev(si.has_og_tags) == "No":
            mktg_insights.append("No Open Graph tags detected — social media previews will fall back to generic look. Adding OG tags improves click-through rates from social shares.")
        if _ev(si.has_robots_txt) == "Found":
            mktg_insights.append("robots.txt is accessible — search engines can find your crawl directives.")
        elif _ev(si.has_robots_txt) == "Not Found":
            mktg_insights.append("robots.txt not found — search engines will crawl everything by default, which may waste crawl budget.")
        if _ev(si.has_sitemap_xml) == "Found":
            mktg_insights.append("XML sitemap detected — helps Google discover and index your pages efficiently.")
        elif _ev(si.has_sitemap_xml) == "Not Found":
            mktg_insights.append("No XML sitemap found — creating one is recommended for proper indexation.")
        if mktg_insights:
            _bullets(doc, mktg_insights)

    # ════════════════════════════════════════════
    # CARD 6: GSC PERFORMANCE (conditional)
    # ════════════════════════════════════════════
    k = facts.kpis
    has_gsc = (_ev(k.clicks) != "\u2014" or _ev(k.impressions) != "\u2014")
    if has_gsc:
        doc.add_heading("Google Search Console Performance", level=1)
        _para(doc, "Month-over-month comparison of how users discover your website through Google search.", size=9, color=GRAY)
        gsc_data = []
        if _ev(k.clicks) != "\u2014":
            gsc_data.append(("Total Clicks", _ev(k.clicks)))
        if _ev(k.impressions) != "\u2014":
            gsc_data.append(("Total Impressions", _ev(k.impressions)))
        if _ev(k.clicks_change) != "\u2014":
            cc = _ev(k.clicks_change)
            gsc_data.append(("Clicks Change", f"{cc}%" if "%" not in cc else cc))
        if _ev(k.impressions_change) != "\u2014":
            ic = _ev(k.impressions_change)
            gsc_data.append(("Impressions Change", f"{ic}%" if "%" not in ic else ic))

        if gsc_data:
            gt = doc.add_table(rows=1 + len(gsc_data), cols=2)
            gt.alignment = WD_TABLE_ALIGNMENT.CENTER
            _cell(gt.rows[0].cells[0], "Metric", bold=True, color=WHITE, size=9)
            _cell(gt.rows[0].cells[1], "Value", bold=True, color=WHITE, size=9)
            _shd(gt.rows[0].cells[0], "0F2747")
            _shd(gt.rows[0].cells[1], "0F2747")
            for ri, (lb, vl) in enumerate(gsc_data):
                _cell(gt.rows[ri+1].cells[0], lb, size=9)
                _cell(gt.rows[ri+1].cells[1], vl, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                if ri % 2 == 0:
                    _shd(gt.rows[ri+1].cells[0], "F8FAFC")
                    _shd(gt.rows[ri+1].cells[1], "F8FAFC")

        _para(doc, "", size=4)
        gsc_insights = []
        if _ev(k.impressions) != "\u2014":
            gsc_insights.append(f"Impressions: {_ev(k.impressions)}")
        if _ev(k.clicks) != "\u2014":
            gsc_insights.append(f"Clicks: {_ev(k.clicks)}")
        if _ev(k.clicks_change) != "\u2014":
            gsc_insights.append(f"Clicks change: {_ev(k.clicks_change)}")
        if _ev(k.impressions_change) != "\u2014":
            gsc_insights.append(f"Impressions change: {_ev(k.impressions_change)}")
        if gsc_insights:
            _para(doc, "GSC metrics:", bold=True, size=10, color=NAVY, before=2)
            _bullets(doc, gsc_insights)

        # GA4 organic traffic inline
        has_ga4 = _ev(k.organic_users) not in ("\u2014", "Access Required")
        if has_ga4:
            _para(doc, "", size=4)
            _para(doc, "Google Analytics 4 — Organic Traffic", bold=True, size=10, color=NAVY)
            ga4_items = []
            if _ev(k.organic_users) != "\u2014":
                ga4_items.append(("Organic Users", _ev(k.organic_users)))
            if _ev(k.sessions) != "\u2014":
                ga4_items.append(("Organic Sessions", _ev(k.sessions)))
            if ga4_items:
                at = doc.add_table(rows=1 + len(ga4_items), cols=2)
                at.alignment = WD_TABLE_ALIGNMENT.CENTER
                _cell(at.rows[0].cells[0], "Metric", bold=True, color=WHITE, size=9)
                _cell(at.rows[0].cells[1], "Value", bold=True, color=WHITE, size=9)
                _shd(at.rows[0].cells[0], "0F2747")
                _shd(at.rows[0].cells[1], "0F2747")
                for ri, (lb, vl) in enumerate(ga4_items):
                    _cell(at.rows[ri+1].cells[0], lb, size=9)
                    _cell(at.rows[ri+1].cells[1], vl, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    if ri % 2 == 0:
                        _shd(at.rows[ri+1].cells[0], "F8FAFC")
                        _shd(at.rows[ri+1].cells[1], "F8FAFC")
    else:
        _access_required(doc, "Google Search Console",
                         "Connect your GSC account to see clicks, impressions, CTR, and average position data.")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 7: GOOGLE TRENDS
    # ════════════════════════════════════════════
    if google_trends_data:
        timeline = google_trends_data.get("timeline_data", [])
        if timeline:
            doc.add_heading("Google Search Interest Over Time", level=1)
            _para(doc, "How search interest in your brand and category keywords has trended over the past 12 months, sourced from Google Trends.", size=9, color=GRAY)
            gt = doc.add_table(rows=1 + len(timeline), cols=2)
            gt.alignment = WD_TABLE_ALIGNMENT.CENTER
            _cell(gt.rows[0].cells[0], "Week", bold=True, color=WHITE, size=8)
            _cell(gt.rows[0].cells[1], "Search Interest", bold=True, color=WHITE, size=8)
            _shd(gt.rows[0].cells[0], "0F2747")
            _shd(gt.rows[0].cells[1], "0F2747")
            for ri, entry in enumerate(timeline):
                _cell(gt.rows[ri+1].cells[0], (entry.get("date") or "")[:20], size=8)
                vals = entry.get("values", [])
                val_str = (vals[0].get("value", "\u2014") if vals else "\u2014")
                _cell(gt.rows[ri+1].cells[1], val_str, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                if ri % 2 == 0:
                    _shd(gt.rows[ri+1].cells[0], "F8FAFC")
                    _shd(gt.rows[ri+1].cells[1], "F8FAFC")
            doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 8: BRAND KEYWORD RANKINGS
    # ════════════════════════════════════════════
    doc.add_heading("Search Visibility & Brand Keyword Rankings", level=1)
    _para(doc, "How your brand and product keywords perform in Google search results.", size=9, color=GRAY)

    if total > 0:
        dist = {"p1": p1, "top3": max(0, top3 - p1), "top10": max(0, top10 - top3),
                "p11_20": sum(1 for r in ranks if 11 <= _evi(r.position) <= 20),
                "beyond": sum(1 for r in ranks if _evi(r.position) > 20)}
        _img(doc, make_distribution_hbar_chart(dist, width=5.0, height=2.5), w=4.8)

    _para(doc, "", size=2)
    rank_insights = []
    if total > 0:
        rank_insights.append(f"Keywords tracked: {total} | #1: {p1} | Top 3: {top3} | Top 10: {top10}")
        rank_insights.append(f"Movement: +{improved} improved, {dropped} dropped, {total - improved - dropped} stable")
    else:
        rank_insights.append("No keyword rankings were collected this period — configure SERP tracking or upload a keyword sheet.")
    _bullets(doc, rank_insights)

    _para(doc, "", size=4)
    _para(doc, "Current keyword positions (top 10 shown):", bold=True, size=10, color=NAVY)

    if total > 0:
        wt = doc.add_table(rows=1 + min(total, 10), cols=4)
        wt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Keyword", "Position", "Change", "Intent"]):
            _cell(wt.rows[0].cells[ci], h, bold=True, color=WHITE, size=8)
            _shd(wt.rows[0].cells[ci], "0F2747")
        for ri, r in enumerate(ranks[:10]):
            chg = _evi(r.change)
            chg_str = f"+{chg}" if chg > 0 else str(chg)
            chg_clr = GREEN if chg > 0 else (RED if chg < 0 else GRAY)
            _cell(wt.rows[ri+1].cells[0], r.keyword[:80], size=8)
            _cell(wt.rows[ri+1].cells[1], _ev(r.position), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(wt.rows[ri+1].cells[2], chg_str, size=8, bold=True, color=chg_clr, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(wt.rows[ri+1].cells[3], r.competition or "\u2014", size=8)
            if ri % 2 == 0:
                for ci in range(4): _shd(wt.rows[ri+1].cells[ci], "F8FAFC")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 9: TOP PERFORMING KEYWORDS
    # ════════════════════════════════════════════
    doc.add_heading("Top Performing Keywords", level=1)
    _para(doc, "These keywords gained the most positions this month \u2014 a sign that your content and SEO efforts are paying off.", size=9, color=GRAY)

    winners = sorted([r for r in ranks if _evi(r.change) > 0], key=lambda r: _evi(r.change), reverse=True)
    if winners:
        wt = doc.add_table(rows=1 + len(winners), cols=4)
        wt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Keyword", "Position", "Gain", "Intent"]):
            _cell(wt.rows[0].cells[ci], h, bold=True, color=WHITE, size=9)
            _shd(wt.rows[0].cells[ci], "0F2747")
        for ri, r in enumerate(winners):
            _cell(wt.rows[ri+1].cells[0], r.keyword[:35], size=9)
            _cell(wt.rows[ri+1].cells[1], _ev(r.position), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(wt.rows[ri+1].cells[2], f"+{_evi(r.change)}", size=9, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(wt.rows[ri+1].cells[3], r.competition or "\u2014", size=9)
            if ri % 2 == 0:
                for ci in range(4): _shd(wt.rows[ri+1].cells[ci], "F8FAFC")

        _para(doc, "", size=4)
        _para(doc, "Biggest gains:", bold=True, size=10, color=NAVY)
        for r in winners[:3]:
            p = doc.add_paragraph(style="List Bullet")
            rn = p.add_run(f"'{r.keyword[:35]}' +{_evi(r.change)} positions (now #{_ev(r.position)})")
            rn.font.size = Pt(9)
            rn.font.name = FN
    else:
        if total > 0:
            _para(doc, "No positive movement this period — review content and internal linking for all tracked keywords.", size=10, color=BODY)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 10: KEYWORDS NEEDING ATTENTION
    # ════════════════════════════════════════════
    doc.add_heading("Keywords Needing Attention", level=1)
    _para(doc, "These keywords dropped in position and need targeted content optimization to recover.", size=9, color=GRAY)

    losers = sorted([r for r in ranks if _evi(r.change) < 0], key=lambda r: _evi(r.change))
    if losers:
        lt = doc.add_table(rows=1 + len(losers), cols=4)
        lt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Keyword", "Position", "Drop", "Intent"]):
            _cell(lt.rows[0].cells[ci], h, bold=True, color=WHITE, size=9)
            _shd(lt.rows[0].cells[ci], "0F2747")
        for ri, r in enumerate(losers):
            _cell(lt.rows[ri+1].cells[0], r.keyword[:35], size=9)
            _cell(lt.rows[ri+1].cells[1], _ev(r.position), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(lt.rows[ri+1].cells[2], str(_evi(r.change)), size=9, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(lt.rows[ri+1].cells[3], r.competition or "\u2014", size=9)
            if ri % 2 == 0:
                for ci in range(4): _shd(lt.rows[ri+1].cells[ci], "F8FAFC")

        _para(doc, "", size=4)
        _para(doc, "Recommended actions:", bold=True, size=10, color=NAVY)
        lose_insights = []
        lose_insights.append(f"{len(losers)} keywords dropped — assign content refresh per keyword.")
        _bullets(doc, lose_insights)
    else:
        if total > 0:
            _para(doc, f"All {total} tracked keywords maintained or improved this period.", size=10, color=BODY)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 10: PRESS COVERAGE
    # ════════════════════════════════════════════
    has_press = bool(facts.press_coverages or any(
        "press" in r.keyword.lower() or "forbes" in r.keyword.lower()
        or "vogue" in r.keyword.lower() for r in ranks))
    if has_press or facts.press_coverages:
        doc.add_heading("Press Coverage & Regional Visibility", level=1)
        _para(doc, "Media mentions and their search performance across tracked markets.", size=9, color=GRAY)

        if facts.press_coverages:
            _para(doc, "The following press coverages rank #1 on Google \u2014 every media feature is easily discoverable:", size=10, color=BODY)
            pt = doc.add_table(rows=1 + len(facts.press_coverages), cols=1)
            pt.alignment = WD_TABLE_ALIGNMENT.CENTER
            _cell(pt.rows[0].cells[0], "Publication / Coverage", bold=True, color=WHITE, size=9)
            _shd(pt.rows[0].cells[0], "0F2747")
            for ri, cov in enumerate(facts.press_coverages):
                _cell(pt.rows[ri+1].cells[0], cov, size=9)
                if ri % 2 == 0:
                    _shd(pt.rows[ri+1].cells[0], "F8FAFC")

            _para(doc, "", size=4)
            _para(doc, "Note:", bold=True, size=10, color=NAVY)
            press_p1 = sum(1 for r in ranks if "press" in r.keyword.lower() and _evi(r.position) == 1)
            press_insights = [
                f"{len(facts.press_coverages)} press coverages tracked. {press_p1} hold #1 position.",
            ]
            if press_p1 > 0:
                press_insights.append("First-page press coverage contributes to brand authority signals.")
            _bullets(doc, press_insights)

        press_ranks = [r for r in ranks if "press" in r.keyword.lower()
                       or "forbes" in r.keyword.lower() or "vogue" in r.keyword.lower()
                       or "news" in r.keyword.lower() or "gq" in r.keyword.lower()
                       or "grazia" in r.keyword.lower() or "cosmopolitan" in r.keyword.lower()
                       or "cntraveller" in r.keyword.lower() or "praveen kenneth" in r.keyword.lower()]
        if press_ranks:
            _para(doc, "", size=4)
            _para(doc, "Article keyword position tracking:", bold=True, size=10, color=NAVY)
            pt2 = doc.add_table(rows=1 + len(press_ranks), cols=4)
            pt2.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, h in enumerate(["Keyword", "Position", "Change", "Region"]):
                _cell(pt2.rows[0].cells[ci], h, bold=True, color=WHITE, size=8)
                _shd(pt2.rows[0].cells[ci], "0F2747")
            for ri, r in enumerate(press_ranks):
                _cell(pt2.rows[ri+1].cells[0], r.keyword[:40], size=8)
                _cell(pt2.rows[ri+1].cells[1], _ev(r.position), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                chg = _evi(r.change)
                _cell(pt2.rows[ri+1].cells[2], f"+{chg}" if chg > 0 else str(chg), size=8, bold=True,
                      color=GREEN if chg > 0 else (RED if chg < 0 else GRAY), align=WD_ALIGN_PARAGRAPH.CENTER)
                _cell(pt2.rows[ri+1].cells[3], "India" if "india" in r.keyword.lower() else "Paris/Global", size=8)
                if ri % 2 == 0:
                    for ci in range(4): _shd(pt2.rows[ri+1].cells[ci], "F8FAFC")

        doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 11: LOCAL SEO
    # ════════════════════════════════════════════
    if gmb_total > 0 or facts.local_seo:
        doc.add_heading("Local SEO \u2014 Google My Business", level=1)
        _para(doc, "How your GMB profile performs for local searches. This directly drives foot traffic to your flagship store and local discovery of your brand.", size=9, color=GRAY)

        ls = facts.local_seo
        if ls and (_ev(ls.review_count) != "\u2014" or _ev(ls.avg_rating) != "\u2014"):
            gmb_info = []
            if _ev(ls.review_count) != "\u2014":
                gmb_info.append(("Reviews", _ev(ls.review_count)))
            if _ev(ls.avg_rating) != "\u2014":
                gmb_info.append(("Avg Rating", _ev(ls.avg_rating)))
            if _ev(ls.gmb_posts) != "\u2014":
                gmb_info.append(("GMB Posts", _ev(ls.gmb_posts)))
            if gmb_info:
                git = doc.add_table(rows=1 + len(gmb_info), cols=2)
                git.alignment = WD_TABLE_ALIGNMENT.CENTER
                _cell(git.rows[0].cells[0], "GMB Profile", bold=True, color=WHITE, size=9)
                _cell(git.rows[0].cells[1], "Status", bold=True, color=WHITE, size=9)
                _shd(git.rows[0].cells[0], "0F2747")
                _shd(git.rows[0].cells[1], "0F2747")
                for ri, (lb, vl) in enumerate(gmb_info):
                    _cell(git.rows[ri+1].cells[0], lb, size=9)
                    _cell(git.rows[ri+1].cells[1], vl, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    if ri % 2 == 0:
                        _shd(git.rows[ri+1].cells[0], "F8FAFC")
                        _shd(git.rows[ri+1].cells[1], "F8FAFC")

        if gmb_total > 0:
            _para(doc, "", size=4)
            _para(doc, f"Your GMB profile is tracked across {gmb_total} local keywords. {gmb_p1} keywords hold #1 in local pack, {gmb_top3} are in the top 3.", size=10, color=BODY)

            max_dates = max(len(k.dates) for k in facts.gmb_keywords) if facts.gmb_keywords else 0
            cols = 2 + max_dates
            gt = doc.add_table(rows=1 + gmb_total, cols=cols)
            gt.alignment = WD_TABLE_ALIGNMENT.CENTER
            _cell(gt.rows[0].cells[0], "Keyword", bold=True, color=WHITE, size=8)
            _cell(gt.rows[0].cells[1], "Type", bold=True, color=WHITE, size=8)
            for di in range(max_dates):
                lbl = facts.gmb_keywords[0].dates[di][-8:] if facts.gmb_keywords and di < len(facts.gmb_keywords[0].dates) else f"D{di+1}"
                _cell(gt.rows[0].cells[2 + di], lbl, bold=True, color=WHITE, size=7)
                _shd(gt.rows[0].cells[2 + di], "0F2747")
            _shd(gt.rows[0].cells[0], "0F2747")
            _shd(gt.rows[0].cells[1], "0F2747")
            for ri, kw in enumerate(facts.gmb_keywords):
                _cell(gt.rows[ri+1].cells[0], kw.keyword[:30], size=8)
                _cell(gt.rows[ri+1].cells[1], "GMB Local", size=7, color=GRAY)
                for di in range(max_dates):
                    pos = kw.positions[di] if di < len(kw.positions) else "\u2014"
                    pv = _evi(pos)
                    pc = GREEN if pv == 1 else (ACCENT if pv <= 3 else (AMBER if pv <= 10 else RED))
                    _cell(gt.rows[ri+1].cells[2 + di], str(pos), size=8, bold=True, color=pc, align=WD_ALIGN_PARAGRAPH.CENTER)
                if ri % 2 == 0:
                    for ci in range(cols): _shd(gt.rows[ri+1].cells[ci], "F8FAFC")

        # Observations
        obs = ls.gmb_observations if ls else []
        if obs:
            _para(doc, "", size=4)
            _para(doc, "Observations:", bold=True, size=10, color=NAVY)
            for ob in obs:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(ob)
                r.font.size = Pt(9)
                r.font.name = FN

        if ls and ls.gmb_content_note:
            _para(doc, "", size=2)
            _para(doc, f"GMB Content Update: {ls.gmb_content_note}", size=9, color=BODY)

        gmb_takeaways = []
        if gmb_p1 > 0:
            gmb_takeaways.append(f"{gmb_p1} local keywords hold the #1 spot in Google Maps.")
        if gmb_top3 > 0:
            gmb_takeaways.append(f"{gmb_top3} local keywords appear in the top 3 Google Maps results.")
        if ls and _evi(ls.review_count) > 0:
            gmb_takeaways.append(f"{_ev(ls.review_count)} reviews at {_ev(ls.avg_rating)} stars on GMB.")

        if gmb_takeaways:
            _para(doc, "", size=4)
            _para(doc, "What this means for your store:", bold=True, size=10, color=NAVY)
            _bullets(doc, gmb_takeaways)

        doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 12: OFF-PAGE SEO ACTIVITIES
    # ════════════════════════════════════════════
    has_offpage = bool(facts.social_bookmarks or facts.image_submissions
                       or facts.video_submissions)
    if has_offpage:
        doc.add_heading("Off-Page SEO Activities", level=1)
        _para(doc, "Link building and content promotion activities completed this month to strengthen your website's authority and search presence beyond on-page optimization.", size=9, color=GRAY)

        _para(doc, "Activities completed this period:", bold=True, size=10, color=NAVY)

        for title, items in [("Social Bookmarking", facts.social_bookmarks),
                              ("Image Submissions", facts.image_submissions),
                              ("Video Submissions", facts.video_submissions)]:
            if items:
                _para(doc, "", size=3)
                _para(doc, f"{title} ({len(items)} submissions)", bold=True, size=11, color=ACCENT)
                st = doc.add_table(rows=1 + len(items), cols=3)
                st.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ci, h in enumerate(["Platform", "Target Content", "Submitted URL"]):
                    _cell(st.rows[0].cells[ci], h, bold=True, color=WHITE, size=7)
                    _shd(st.rows[0].cells[ci], "0F2747")
                for ri, (site, target, url) in enumerate(items):
                    _cell(st.rows[ri+1].cells[0], site[:25], size=7)
                    _cell(st.rows[ri+1].cells[1], target[:40], size=7)
                    _cell(st.rows[ri+1].cells[2], url[:40], size=6, color=GRAY)
                    if ri % 2 == 0:
                        for ci in range(3): _shd(st.rows[ri+1].cells[ci], "F8FAFC")

        sbm = len(facts.social_bookmarks)
        imgm = len(facts.image_submissions)
        vidm = len(facts.video_submissions)
        total_offpage = sbm + imgm + vidm
        _para(doc, "", size=4)
        _para(doc, f"Activity summary: {total_offpage} total submissions completed ({sbm} social bookmarks, {imgm} image submissions, {vidm} video submissions).", size=10, color=BODY)

        doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 13: SEO ACTIVITIES COMPLETED
    # ════════════════════════════════════════════
    if facts.seo_activities_completed:
        doc.add_heading("SEO Activities Performed", level=1)
        _para(doc, "A comprehensive summary of all SEO work completed this period \u2014 covering on-page optimization, off-page authority building, and technical improvements.", size=9, color=GRAY)

        _para(doc, "", size=2)
        onpage = [a for a in facts.seo_activities_completed if any(k in a.lower() for k in ["blog", "content", "keyword", "on-page", "schema", "article"])]
        offpage_items = [a for a in facts.seo_activities_completed if any(k in a.lower() for k in ["link building", "guest", "outreach", "bookmarking", "submission", "quora", "web 2.0"])]
        tech_activities = [a for a in facts.seo_activities_completed if any(k in a.lower() for k in ["technical", "gmb", "schema", "image schema"])]

        categories = [
            ("On-Page & Content", onpage),
            ("Off-Page & Authority", offpage_items),
            ("Technical & Local", tech_activities),
        ]
        for cat_name, items in categories:
            if items:
                _para(doc, cat_name, bold=True, size=11, color=NAVY)
                for act in items:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.space_after = Pt(1)
                    r = p.add_run(act)
                    r.font.size = Pt(9)
                    r.font.name = FN
                _para(doc, "", size=2)

        _para(doc, f"Total: {len(facts.seo_activities_completed)} activities completed this period.", bold=True, size=10, color=BODY)
        doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 14: CORE WEB VITALS
    # ════════════════════════════════════════════
    if cm > 0 or cd > 0:
        doc.add_heading("Core Web Vitals", level=1)
        _para(doc, "Google's user experience metrics \u2014 directly impact search rankings and how users experience your website.", size=9, color=GRAY)

        if cm > 0 and cd > 0:
            _img(doc, make_psi_column_chart(cd, cm, width=3.5, height=2.5), w=3.5)

        ct = doc.add_table(rows=3, cols=2)
        ct.alignment = WD_TABLE_ALIGNMENT.CENTER
        _cell(ct.rows[0].cells[0], "Device", bold=True, color=WHITE, size=9)
        _cell(ct.rows[0].cells[1], "Score", bold=True, color=WHITE, size=9)
        _shd(ct.rows[0].cells[0], "0F2747")
        _shd(ct.rows[0].cells[1], "0F2747")
        ds, dc = _cstatus(cd) if cd > 0 else ("N/A", GRAY)
        ms, mc = _cstatus(cm) if cm > 0 else ("N/A", GRAY)
        _cell(ct.rows[1].cells[0], "Desktop", bold=True, size=10)
        _cell(ct.rows[1].cells[1], f"{cd}/100 ({ds})", size=10, bold=True, color=dc)
        _shd(ct.rows[1].cells[0], "F8FAFC")
        _shd(ct.rows[1].cells[1], "F8FAFC")
        _cell(ct.rows[2].cells[0], "Mobile", bold=True, size=10)
        _cell(ct.rows[2].cells[1], f"{cm}/100 ({ms})", size=10, bold=True, color=mc)

        _para(doc, "", size=4)
        cwv_insights = []
        lcp_val = _ev(facts.cwv.lcp_seconds)
        inp_val = _ev(facts.cwv.inp_ms)
        cls_val = _ev(facts.cwv.cls_score)
        cwv_data = []
        if cd > 0:
            cwv_data.append(f"Desktop: {cd}/100")
        if cm > 0:
            cwv_data.append(f"Mobile: {cm}/100")
        if lcp_val != "\u2014":
            cwv_data.append(f"LCP: {lcp_val}s")
        if inp_val != "\u2014":
            cwv_data.append(f"INP: {inp_val}ms")
        if cls_val != "\u2014":
            cwv_data.append(f"CLS: {cls_val}")
        if cwv_data:
            cwv_insights.append(" | ".join(cwv_data))
        else:
            cwv_insights.append("No CWV data collected — connect PageSpeed API to see lab metrics.")
        _bullets(doc, cwv_insights)
        doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 15: TECHNICAL SEO HEALTH
    # ════════════════════════════════════════════
    if th > 0 or t.issues_list:
        doc.add_heading("Technical SEO Health", level=1)
        _para(doc, "A technical audit identifies issues that may prevent search engines from properly crawling, indexing, and ranking your pages.", size=9, color=GRAY)

        if th > 0:
            sc = GREEN if th >= 80 else (AMBER if th >= 50 else RED)
            _para(doc, f"Health Score: {th}/100", bold=True, color=sc, size=13)

        summary = [("Pages Audited", _ev(t.pages_audited)),
                   ("Missing H1", _ev(t.missing_h1)),
                   ("Missing Meta", _ev(t.missing_meta)),
                   ("Missing Alt Tags", _ev(t.missing_alt)),
                   ("Thin Pages", _ev(t.thin_pages))]
        st = doc.add_table(rows=1 + len(summary), cols=2)
        st.alignment = WD_TABLE_ALIGNMENT.CENTER
        _cell(st.rows[0].cells[0], "Category", bold=True, color=WHITE, size=9)
        _cell(st.rows[0].cells[1], "Count", bold=True, color=WHITE, size=9)
        _shd(st.rows[0].cells[0], "0F2747")
        _shd(st.rows[0].cells[1], "0F2747")
        for ri, (lb, vl) in enumerate(summary):
            _cell(st.rows[ri+1].cells[0], lb, size=9)
            _cell(st.rows[ri+1].cells[1], vl, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2 == 0:
                _shd(st.rows[ri+1].cells[0], "F8FAFC")
                _shd(st.rows[ri+1].cells[1], "F8FAFC")

        _para(doc, "", size=4)
        tech_insights = []
        mi_h1 = _evi(t.missing_h1)
        mi_m = _evi(t.missing_meta)
        mi_a = _evi(t.missing_alt)
        if mi_h1 > 0:
            tech_insights.append(f"{mi_h1} pages are missing H1 headings \u2014 these are important for both SEO and accessibility.")
        if mi_m > 0:
            tech_insights.append(f"{mi_m} pages are missing meta descriptions \u2014 adding them improves click-through rates from search results.")
        if mi_a > 0:
            tech_insights.append(f"{mi_a} images lack alt text \u2014 this is a quick fix that improves accessibility and image search visibility.")
        _bullets(doc, tech_insights)

        if t.issues_list:
            _para(doc, "", size=4)
            _para(doc, f"Issues Detected ({len(t.issues_list)} found):", bold=True, size=10, color=NAVY)
            it = doc.add_table(rows=1 + len(t.issues_list), cols=3)
            it.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, h in enumerate(["Severity", "Issue", "Page"]):
                _cell(it.rows[0].cells[ci], h, bold=True, color=WHITE, size=8)
                _shd(it.rows[0].cells[ci], "0F2747")
            for ri, iss in enumerate(t.issues_list):
                sc = _sev_color(iss.severity)
                _cell(it.rows[ri+1].cells[0], iss.severity.upper(), bold=True, color=sc, size=8)
                _cell(it.rows[ri+1].cells[1], iss.issue_text[:55], size=8)
                _cell(it.rows[ri+1].cells[2], iss.page, size=7, color=GRAY)
                if ri % 2 == 0:
                    for ci in range(3): _shd(it.rows[ri+1].cells[ci], "F8FAFC")

            crit_count = sum(1 for iss in t.issues_list if iss.severity.lower() == "critical")
            if crit_count > 0:
                _para(doc, "", size=2)
                _para(doc, f"Critical issues ({crit_count}) will be prioritized this week \u2014 these directly impact how Google crawls and indexes your site.", size=9, color=BODY)

        doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 16: RECOMMENDATIONS
    # ════════════════════════════════════════════
    doc.add_heading("Recommendations & Strategy for Next Month", level=1)
    _para(doc, "Prioritised actions for the coming period, organised by urgency and impact.", size=9, color=GRAY)

    if facts.way_forward:
        _para(doc, "", size=2)
        _para(doc, "Planned activities for next month:", bold=True, size=11, color=NAVY)
        for wf in facts.way_forward:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(wf)
            r.font.size = Pt(9)
            r.font.name = FN

    _para(doc, "", size=4)
    _para(doc, "Priority breakdown for execution:", bold=True, size=11, color=NAVY)

    if facts.action_plan:
        for p_lbl, p_filter, limit in [("P0 \u2014 Immediate (This Week)", lambda a: a.priority == "P0", 4),
                                        ("P1 \u2014 High Priority (Next Week)", lambda a: a.priority == "P1", 6),
                                        ("P2 \u2014 Medium (This Sprint)", lambda a: a.priority == "P2", 6)]:
            items = [a for a in facts.action_plan if p_filter(a)]
            if items:
                _para(doc, p_lbl, bold=True, size=10, color=ACCENT)
                for a in items[:limit]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.space_after = Pt(1)
                    r = p.add_run(f"[{a.team}] {a.task[:55]}")
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.name = FN
                    p.add_run(f"  ETA: {a.eta}").font.size = Pt(8)
                _para(doc, "", size=2)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 17: ACTION PLAN BY TEAM
    # ════════════════════════════════════════════
    if facts.action_plan:
        doc.add_heading("Action Plan by Team", level=1)
        _para(doc, "Who owns what \u2014 every action item assigned to a specific team with priority level.", size=9, color=GRAY)

        teams = ["SEO", "Dev", "Content", "Local", "Off-Page"]
        team_desc = {
            "SEO": "Keyword optimisation, technical fixes, site structure",
            "Dev": "CWV improvement, site performance, technical implementation",
            "Content": "Content creation, keyword recovery, blog strategy",
            "Local": "GMB management, local listings, review responses",
            "Off-Page": "Link building, outreach, social bookmarking, PR SEO",
        }
        team_items = {t: [a for a in facts.action_plan if a.team == t] for t in teams}
        tt = doc.add_table(rows=1 + len(teams), cols=3)
        tt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Team", "Tasks", "Top Priority"]):
            _cell(tt.rows[0].cells[ci], h, bold=True, color=WHITE, size=9)
            _shd(tt.rows[0].cells[ci], "0F2747")
        for ri, team in enumerate(teams):
            items = team_items[team]
            top_p = min(items, key=lambda a: a.priority).priority if items else "\u2014"
            _cell(tt.rows[ri+1].cells[0], team, bold=True, size=9)
            _cell(tt.rows[ri+1].cells[1], str(len(items)), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(tt.rows[ri+1].cells[2], top_p, size=9, bold=True,
                  color=RED if top_p in ("P0", "P1") else AMBER)
            if ri % 2 == 0:
                for ci in range(3): _shd(tt.rows[ri+1].cells[ci], "F8FAFC")

        _para(doc, "", size=4)
        _para(doc, "Team responsibilities:", bold=True, size=10, color=NAVY)
        for team in teams:
            desc = team_desc.get(team, "")
            _para(doc, f"\u2022 {team}: {desc}", size=9, color=BODY, after=1)

        doc.add_page_break()

    # ════════════════════════════════════════════
    # CARD 18: CLOSING
    # ════════════════════════════════════════════
    doc.add_heading("Summary & Next Steps", level=1)
    total_ranked = len([r for r in ranks if _evi(r.position) > 0])
    summary_parts = []
    if total_ranked > 0:
        summary_parts.append(f"{total_ranked} tracked keywords")
    if p1 > 0:
        summary_parts.append(f"{p1} at #1")
    if top10 > 0:
        summary_parts.append(f"{top10} in top 10")
    if _evi(t.pages_audited) > 0:
        summary_parts.append(f"{_evi(t.pages_audited)} pages audited")
    summary_text = f"This audit covered {', '.join(summary_parts)}." if summary_parts else "Audit complete."
    _para(doc, summary_text, size=10, color=BODY)

    _para(doc, "", size=4)
    _para(doc, "Priorities for the upcoming period:", bold=True, size=11, color=NAVY)

    priorities = []
    if dropped > 0:
        priorities.append(f"Recover rankings for {dropped} dropped keywords")
    if cm > 0 and cm < 80:
        lcp_val = _ev(facts.cwv.lcp_seconds)
        lcp_target = f" (LCP at {lcp_val})" if lcp_val != "\u2014" else ""
        priorities.append(f"Improve mobile CWV from {cm}/100{lcp_target}")
    if any(iss.severity.lower() == "critical" for iss in t.issues_list):
        crit_issues = [iss for iss in t.issues_list if iss.severity.lower() == "critical"]
        issue_summary = "; ".join(i.issue_text[:40] for i in crit_issues[:3])
        priorities.append(f"Fix critical issues: {issue_summary}")
    if gmb_total > 0:
        priorities.append(f"Continue GMB management ({gmb_total} keywords tracked)")
    if facts.way_forward:
        for wf in facts.way_forward[:3]:
            priorities.append(wf.action[:60])

    for p in priorities:
        pp = doc.add_paragraph(style="List Bullet")
        pp.paragraph_format.space_after = Pt(2)
        r = pp.add_run(p)
        r.font.size = Pt(10)
        r.font.name = FN

    _para(doc, "", size=8)
    _para(doc, f"This report was prepared exclusively for {m.client_name or 'the client'} by {m.agency_name or 'the SEO team'}. For questions or a walkthrough of any section, please reach out to your account manager.", color=GRAY, size=8)

    for sec in doc.sections:
        ftr = sec.footer
        ftr.is_linked_to_previous = False
        p = ftr.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("\u2500" * 50)
        r.font.size = Pt(5)
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
        p2 = ftr.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = [m.agency_name] if m.agency_name else []
        parts.append("CONFIDENTIAL")
        if m.client_name: parts.append(m.client_name)
        r = p2.add_run("  |  ".join(parts))
        r.font.size = Pt(7)
        r.font.color.rgb = GRAY
        r.font.name = FN

    doc.save(output_path)
    return output_path
