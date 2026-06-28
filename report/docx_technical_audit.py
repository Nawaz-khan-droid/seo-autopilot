"""Technical SEO Audit — Focused DOCX (no monthly narrative fluff).

Covers: cover, executive summary (optional), technical dashboard, Core Web Vitals,
core SEO metrics, internal link health, backlink profile, issues log, recommendations.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from report.facts import ReportFacts
from report.charts import make_psi_column_chart

logger = logging.getLogger(__name__)

NAVY = RGBColor(0x0F, 0x27, 0x47)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
DARK = RGBColor(0x1E, 0x29, 0x3B)
BODY = RGBColor(0x33, 0x41, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
RED = RGBColor(0xDC, 0x26, 0x26)
TEAL = RGBColor(0x14, 0xB8, 0xA6)
GRAY = RGBColor(0x94, 0xA3, 0xB8)

FN = "Calibri"


def _shd(cell, c):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _cell(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FN
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align is not None: p.alignment = align

def _para(doc, text, bold=False, color=None, size=10, align=None, before=0, after=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FN
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p

def _bullets(doc, items, size=10):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(item)
        r.font.size = Pt(size)
        r.font.name = FN

def _ev(v):
    if v is None: return "\u2014"
    if hasattr(v, "is_available") and not v.is_available: return "\u2014"
    val = v.value if hasattr(v, "value") else v
    return str(val) if val is not None else "\u2014"

def _evi(v):
    try: return int(float(_ev(v)))
    except: return 0

def _img(doc, img_bytes, w=5.0):
    if not img_bytes or len(img_bytes) < 100: return
    s = io.BytesIO(img_bytes)
    doc.add_picture(s, width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def _cstatus(s):
    if s >= 90: return "Good", GREEN
    if s >= 50: return "Needs Work", AMBER
    return "Poor", RED

def _sev_color(s):
    s = s.lower().strip()
    if s == "critical": return RED
    if s in ("high", "error"): return AMBER
    if s in ("medium", "warning"): return ACCENT
    return GRAY


def build_technical_audit(facts: ReportFacts, output_path: str = "",
                          page_preview_bytes: bytes | None = None) -> str:
    m = facts.metadata
    if not output_path:
        sn = (m.client_name or "Client").replace(" ", "_")
        output_path = f"{sn}_Technical_Audit.docx"

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = FN
    style.font.size = Pt(10)
    style.font.color.rgb = BODY

    for lvl, (sz, clr, sb) in {1: (16, NAVY, 4), 2: (13, ACCENT, 3), 3: (11, DARK, 2)}.items():
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = FN
        hs.font.size = Pt(sz)
        hs.font.color.rgb = clr
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(sb)
        hs.paragraph_format.space_after = Pt(2)

    t = facts.technical
    th = _evi(t.health_score)
    cm = _evi(facts.cwv.mobile_score)
    cd = _evi(facts.cwv.desktop_score)

    # ── COVER ──
    for _ in range(5): doc.add_paragraph()
    _para(doc, "TECHNICAL SEO AUDIT", bold=True, color=NAVY, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "\u2500" * 36, color=ACCENT, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "", size=4)
    if m.client_name:
        _para(doc, m.client_name, bold=True, color=NAVY, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, m.report_month or datetime.now().strftime("%B %Y"), color=BODY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    gen = m.generated_at or datetime.now().strftime("%B %d, %Y")
    _para(doc, f"Generated: {gen}", color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ── PAGE PREVIEW (SERP screenshot) ──
    if page_preview_bytes:
        doc.add_heading("Page Preview", level=1)
        _para(doc, "Rendered view of the analysed page at audit time.", size=9, color=GRAY)
        _img(doc, page_preview_bytes, w=5.5)
        doc.add_page_break()

    # ── EXECUTIVE SUMMARY (optional) ──
    if facts.executive_narrative:
        doc.add_heading("Executive Summary", level=1)
        _para(doc, facts.executive_narrative, size=10, color=BODY)
        _para(doc, "", size=4)

    # ── TECHNICAL DASHBOARD ──
    doc.add_heading("Technical Dashboard", level=1)
    _para(doc, "At-a-glance technical health metrics from the automated crawl.", size=9, color=GRAY)

    kpi_items = [
        ("Tech Health", f"{th}/100", GREEN if th >= 80 else (AMBER if th >= 50 else RED), "16A34A" if th >= 80 else ("F59E0B" if th >= 50 else "DC2626")),
        ("Pages Audited", _ev(t.pages_audited), NAVY, "0F2747"),
        ("Total Issues", _ev(t.total_issues), RED if _evi(t.total_issues) > 0 else GREEN, "DC2626" if _evi(t.total_issues) > 0 else "16A34A"),
    ]
    kt = doc.add_table(rows=2, cols=len(kpi_items))
    kt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (lb, vl, tc, bg) in enumerate(kpi_items):
        _cell(kt.rows[0].cells[ci], lb, bold=True, color=WHITE, size=8)
        _shd(kt.rows[0].cells[ci], bg)
        _cell(kt.rows[1].cells[ci], vl, bold=True, color=tc, size=15)
        _shd(kt.rows[1].cells[ci], "F8FAFC")

    _para(doc, "", size=4)

    # Missing tags summary
    summary = [("Missing H1", _ev(t.missing_h1)), ("Missing Meta Descriptions", _ev(t.missing_meta)),
               ("Missing Alt Tags", _ev(t.missing_alt)), ("Thin Pages", _ev(t.thin_pages)),
               ("Schema Markup", _ev(t.has_schema))]
    st = doc.add_table(rows=1 + len(summary), cols=2)
    st.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell(st.rows[0].cells[0], "Issue Category", bold=True, color=WHITE, size=9)
    _cell(st.rows[0].cells[1], "Count", bold=True, color=WHITE, size=9)
    _shd(st.rows[0].cells[0], "0F2747")
    _shd(st.rows[0].cells[1], "0F2747")
    for ri, (lb, vl) in enumerate(summary):
        _cell(st.rows[ri+1].cells[0], lb, size=9)
        _cell(st.rows[ri+1].cells[1], vl, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        if ri % 2 == 0:
            _shd(st.rows[ri+1].cells[0], "F8FAFC")
            _shd(st.rows[ri+1].cells[1], "F8FAFC")

    doc.add_page_break()

    # ── CORE WEB VITALS ──
    if cm > 0 or cd > 0:
        doc.add_heading("Core Web Vitals", level=1)
        _para(doc, "Google user experience metrics — directly impact search rankings.", size=9, color=GRAY)

        if cm > 0 and cd > 0:
            _img(doc, make_psi_column_chart(cd, cm, width=3.5, height=2.5), w=3.5)

        ct = doc.add_table(rows=3, cols=3)
        ct.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Device", "Score", "Status"]):
            _cell(ct.rows[0].cells[ci], h, bold=True, color=WHITE, size=9)
            _shd(ct.rows[0].cells[ci], "0F2747")
        ds, dc = _cstatus(cd) if cd > 0 else ("N/A", GRAY)
        ms, mc = _cstatus(cm) if cm > 0 else ("N/A", GRAY)
        for ri, (dev, sc, sc_color, st_lbl) in enumerate([
            ("Desktop", f"{cd}/100" if cd > 0 else "—", dc, ds),
            ("Mobile", f"{cm}/100" if cm > 0 else "—", mc, ms),
        ]):
            _cell(ct.rows[ri+1].cells[0], dev, bold=True, size=10)
            _cell(ct.rows[ri+1].cells[1], sc, size=10, bold=True, color=sc_color)
            _cell(ct.rows[ri+1].cells[2], st_lbl, size=9, color=sc_color)
            _shd(ct.rows[ri+1].cells[0], "F8FAFC")
            _shd(ct.rows[ri+1].cells[1], "F8FAFC")
            _shd(ct.rows[ri+1].cells[2], "F8FAFC")

        cwv_detail = []
        if facts.cwv.lcp_seconds.is_available:
            cwv_detail.append(f"LCP: {_ev(facts.cwv.lcp_seconds)}s")
        if facts.cwv.inp_ms.is_available:
            cwv_detail.append(f"INP: {_ev(facts.cwv.inp_ms)}ms")
        if facts.cwv.cls_score.is_available:
            cwv_detail.append(f"CLS: {_ev(facts.cwv.cls_score)}")
        if cwv_detail:
            _para(doc, "", size=4)
            _para(doc, "Lab Metrics: " + " | ".join(cwv_detail), size=9, color=BODY)

        doc.add_page_break()

    # ── CORE SEO METRICS ──
    si = facts.site_info
    if si.title_tag or si.meta_description or si.h1_count > 0:
        doc.add_heading("Core SEO Metrics", level=1)
        _para(doc, "Foundational on-page elements parsed from the rendered DOM.", size=9, color=GRAY)

        items = []
        if si.title_tag:
            tl = len(si.title_tag)
            items.append(("Title", f"{si.title_tag[:60]} ({tl} chars)", "Good" if 30 <= tl <= 60 else "Needs Work"))
        if si.meta_description:
            dl = len(si.meta_description)
            items.append(("Meta Desc", f"{si.meta_description[:80]} ({dl} chars)", "Good" if 120 <= dl <= 160 else "Needs Work"))
        items.append(("H1 Tags", f"{si.h1_count} found", "Good" if si.h1_count == 1 else "Warning" if si.h1_count > 1 else "Error"))
        if si.h1_texts:
            items.append(("H1 Text", si.h1_texts[0][:60], ""))
        if si.word_count:
            items.append(("Word Count", f"{si.word_count:,}", "Good" if si.word_count >= 300 else "Thin"))

        ct2 = doc.add_table(rows=1 + len(items), cols=3)
        ct2.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Element", "Value", "Status"]):
            _cell(ct2.rows[0].cells[ci], h, bold=True, color=WHITE, size=9)
            _shd(ct2.rows[0].cells[ci], "0F2747")
        for ri, (el, val, st) in enumerate(items):
            _cell(ct2.rows[ri+1].cells[0], el, bold=True, size=9)
            _cell(ct2.rows[ri+1].cells[1], val, size=8)
            sc = GREEN if "Good" in st else (AMBER if "Warning" in st or "Needs" in st else (RED if "Error" in st else GRAY))
            _cell(ct2.rows[ri+1].cells[2], st if st else "\u2014", bold=True, color=sc, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2 == 0:
                for ci in range(3): _shd(ct2.rows[ri+1].cells[ci], "F8FAFC")

        doc.add_page_break()

    # ── INTERNAL LINK HEALTH ──
    bl = facts.backlinks
    il_total = _evi(bl.onpage_total_links)
    if il_total > 0:
        doc.add_heading("Internal Link Health", level=1)
        _para(doc, "Link distribution and health from the initial page crawl.", size=9, color=GRAY)

        il_int = _evi(bl.onpage_internal_links)
        il_ext = _evi(bl.onpage_external_links)
        ilt = doc.add_table(rows=2, cols=4)
        ilt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, (lb, vl) in enumerate([("Total", str(il_total)), ("Internal", str(il_int)), ("External", str(il_ext)), ("Healthy", str(il_total - _evi(t.total_issues)))]):
            _cell(ilt.rows[0].cells[ci], lb, bold=True, color=WHITE, size=9)
            _shd(ilt.rows[0].cells[ci], "0F2747")
            _cell(ilt.rows[1].cells[ci], vl, bold=True, color=ACCENT, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shd(ilt.rows[1].cells[ci], "F8FAFC")

        broken_issues = [i for i in t.issues_list if "Broken link" in i.issue_text or "Redirect" in i.issue_text]
        if broken_issues:
            _para(doc, "", size=4)
            _para(doc, f"Broken/Redirect URLs ({len(broken_issues)} found):", bold=True, size=10, color=NAVY)
            for iss in broken_issues[:8]:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(f"[{iss.severity}] {iss.issue_text}: {iss.page}")
                r.font.size = Pt(8)
                r.font.name = FN

        doc.add_page_break()

    # ── BACKLINK PROFILE ──
    has_bl = _ev(bl.total_backlinks) != "\u2014"
    has_opr_rating = _ev(bl.domain_rating) != "\u2014"
    awaiting_bl = bl.status == "AWAITING_DATA"
    if has_bl or awaiting_bl or has_opr_rating:
        doc.add_heading("Backlink Profile", level=1)
        _para(doc, "External backlink data from automated checks.", size=9, color=GRAY)

    if has_bl:
        bli = [("Total Backlinks", _ev(bl.total_backlinks))]
        if _ev(bl.ref_domains) != "\u2014": bli.append(("Referring Domains", _ev(bl.ref_domains)))
        if _ev(bl.dofollow_count) != "\u2014": bli.append(("Follow Links", _ev(bl.dofollow_count)))
        if _ev(bl.nofollow_count) != "\u2014": bli.append(("Nofollow Links", _ev(bl.nofollow_count)))
        if _ev(bl.domain_rating) != "\u2014": bli.append(("Domain Rating", _ev(bl.domain_rating)))
        bt = doc.add_table(rows=1 + len(bli), cols=2)
        bt.alignment = WD_TABLE_ALIGNMENT.CENTER
        _cell(bt.rows[0].cells[0], "Metric", bold=True, color=WHITE, size=9)
        _cell(bt.rows[0].cells[1], "Value", bold=True, color=WHITE, size=9)
        _shd(bt.rows[0].cells[0], "0F2747")
        _shd(bt.rows[0].cells[1], "0F2747")
        for ri, (lb, vl) in enumerate(bli):
            _cell(bt.rows[ri+1].cells[0], lb, size=9)
            _cell(bt.rows[ri+1].cells[1], vl, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2 == 0:
                _shd(bt.rows[ri+1].cells[0], "F8FAFC")
                _shd(bt.rows[ri+1].cells[1], "F8FAFC")

    elif has_opr_rating:
        _para(doc, "", size=4)
        opr_table = doc.add_table(rows=1, cols=2)
        opr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _cell(opr_table.rows[0].cells[0], "Domain Authority (OpenPageRank)", bold=True, size=9)
        _cell(opr_table.rows[0].cells[1], _ev(bl.domain_rating), size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(opr_table.rows[0].cells[0], "F8FAFC")
        _shd(opr_table.rows[0].cells[1], "F8FAFC")
        _para(doc, "", size=4)
        _para(doc, "Upload a backlink CSV from Ahrefs, Semrush, or Moz for a full profile including backlink counts, referring domains, and follow/nofollow breakdown.", size=8, color=GRAY)

    elif awaiting_bl:
        _para(doc, "", size=4)
        ct = doc.add_table(rows=1, cols=1)
        ct.alignment = WD_TABLE_ALIGNMENT.CENTER
        _shd(ct.rows[0].cells[0], "FFF8E1")
        _cell(ct.rows[0].cells[0], "", size=4)
        p = ct.rows[0].cells[0].paragraphs[-1]
        r = p.add_run("Action Required: Complete Your Backlink Profile")
        r.bold = True; r.font.size = Pt(10); r.font.name = FN
        r.font.color.rgb = RGBColor(0x8D, 0x6E, 0x00)
        p2 = ct.rows[0].cells[0].add_paragraph()
        r2 = p2.add_run(
            "Backlink data was not detected during this automated audit. "
            "To see your full backlink profile, export a CSV from your preferred "
            "tracking tool (Ahrefs, Semrush, Moz, or Ubersuggest) and upload it "
            "via the Data Uploads dashboard. Required columns: total_backlinks, "
            "ref_domains, dofollow, nofollow, domain_rating, source."
        )
        r2.font.size = Pt(9); r2.font.name = FN
        r2.font.color.rgb = RGBColor(0x8D, 0x6E, 0x00)

    if has_bl or awaiting_bl:
        doc.add_page_break()

    # ── ISSUES LOG ──
    if t.issues_list:
        doc.add_heading("Issues Log", level=1)
        _para(doc, f"All detected issues ({len(t.issues_list)} total).", size=9, color=GRAY)

        it = doc.add_table(rows=1 + len(t.issues_list), cols=3)
        it.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Severity", "Issue", "Page"]):
            _cell(it.rows[0].cells[ci], h, bold=True, color=WHITE, size=8)
            _shd(it.rows[0].cells[ci], "0F2747")
        for ri, iss in enumerate(t.issues_list):
            sc = _sev_color(iss.severity)
            _cell(it.rows[ri+1].cells[0], iss.severity.upper(), bold=True, color=sc, size=8)
            _cell(it.rows[ri+1].cells[1], iss.issue_text[:55], size=8)
            _cell(it.rows[ri+1].cells[2], iss.page, size=7, color=GRAY)
            if ri % 2 == 0:
                for ci in range(3): _shd(it.rows[ri+1].cells[ci], "F8FAFC")

        doc.add_page_break()

    # ── RECOMMENDATIONS ──
    if facts.action_plan:
        doc.add_heading("Recommendations", level=1)
        total_actions = len(facts.action_plan)
        crit_count = len([a for a in facts.action_plan if a.priority == "P0"])
        high_count = len([a for a in facts.action_plan if a.priority == "P1"])
        desc_parts = []
        if total_actions:
            desc_parts.append(f"{total_actions} recommended fixes")
        if crit_count:
            desc_parts.append(f"{crit_count} critical")
        if high_count:
            desc_parts.append(f"{high_count} high priority")
        _para(doc, " \u2014 ".join(desc_parts) + "." if desc_parts else "", size=9, color=GRAY)
        for p_lbl, p_filter, limit in [("Critical (This Week)", lambda a: a.priority == "P0", 5),
                                        ("High Priority (Next Week)", lambda a: a.priority == "P1", 6),
                                        ("Medium (This Sprint)", lambda a: a.priority == "P2", 6)]:
            items = [a for a in facts.action_plan if p_filter(a)]
            if items:
                _para(doc, p_lbl, bold=True, size=10, color=ACCENT)
                for a in items[:limit]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.space_after = Pt(1)
                    r = p.add_run(f"[{a.team}] {a.task[:60]}")
                    r.font.size = Pt(9)
                    r.font.name = FN
                    if a.eta:
                        p.add_run(f"  ETA: {a.eta}").font.size = Pt(8)
                _para(doc, "", size=2)

    # ── FOOTER ──
    for sec in doc.sections:
        ftr = sec.footer
        ftr.is_linked_to_previous = False
        p = ftr.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("\u2500" * 50)
        r.font.size = Pt(5)
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
        p2 = ftr.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = [m.agency_name] if m.agency_name else []
        parts.append("Technical SEO Audit — CONFIDENTIAL")
        if m.client_name: parts.append(m.client_name)
        r = p2.add_run("  |  ".join(parts))
        r.font.size = Pt(7)
        r.font.color.rgb = GRAY
        r.font.name = FN

    doc.save(output_path)
    return output_path
