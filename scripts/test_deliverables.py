"""Integration test: build mock ReportFacts → generate all 3 DOCX + verify content."""
import sys, os, re
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))

from docx import Document
from report.facts import (
    ReportFacts, ReportMetadata, RankingRow, ActionItem, TechnicalIssue,
    TechnicalData, CWVData, BacklinkData, SiteInfoData, Kpidata,
)
from report.evidence import Evidence
from report.docx_technical_audit import build_technical_audit
from report.docx_report import build_clients_report
from report.docx_action_plan import build_action_plan_docx

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_test_output")
os.makedirs(OUT, exist_ok=True)

def make_facts():
    f = ReportFacts(
        metadata=ReportMetadata(
            agency_name="Test Agency",
            client_name="Test Client Inc.",
            report_month="June 2026",
            generated_at="2026-06-13T12:00:00",
        ),
        cwv=CWVData(
            mobile_score=Evidence.verified("62", "PSI mobile"),
            desktop_score=Evidence.verified("85", "PSI desktop"),
            lcp_seconds=Evidence.verified("3.2", "PSI mobile"),
            cls_score=Evidence.verified("0.15", "PSI mobile"),
            inp_ms=Evidence.verified("180", "PSI mobile"),
        ),
        technical=TechnicalData(
            health_score=Evidence.verified("72", "SEO Analyzer"),
            pages_audited=Evidence.verified("5", "Crawl"),
            total_issues=Evidence.verified("3", "Crawl"),
            missing_h1=Evidence.verified("2", "Crawl"),
            missing_meta=Evidence.verified("1", "Crawl"),
            missing_alt=Evidence.verified("4", "Crawl"),
            thin_pages=Evidence.verified("1", "Crawl"),
            issues_list=[
                TechnicalIssue(page="https://example.com/", issue_text="Missing H1 tag", severity="High"),
                TechnicalIssue(page="https://example.com/about", issue_text="Broken link (404)", severity="Critical"),
                TechnicalIssue(page="https://example.com/contact", issue_text="Thin content (<200 words)", severity="Medium"),
            ],
        ),
        backlinks=BacklinkData(
            total_backlinks=Evidence.verified("156", "Ahrefs"),
            ref_domains=Evidence.verified("42", "Ahrefs"),
            dofollow_count=Evidence.verified("98", "Ahrefs"),
            nofollow_count=Evidence.verified("58", "Ahrefs"),
            onpage_total_links=Evidence.verified("47", "Playwright"),
            onpage_internal_links=Evidence.verified("32", "Playwright"),
            onpage_external_links=Evidence.verified("15", "Playwright"),
        ),
        site_info=SiteInfoData(
            title_tag="Test Client Inc. - Luxury Products & Services",
            meta_description="Discover premium products from Test Client Inc. with free shipping worldwide.",
            h1_count=1,
            h1_texts=["Welcome to Test Client Inc."],
            word_count=650,
            has_og_tags=Evidence.verified("Yes", "Playwright DOM"),
            has_robots_txt=Evidence.verified("Found", "HTTP HEAD"),
            has_sitemap_xml=Evidence.verified("Found", "HTTP HEAD"),
        ),
        rankings=[
            RankingRow(keyword="luxury perfume brand india", position=Evidence.verified("2", "Sheet"), change=Evidence.verified("0", "Sheet"), competition="high"),
            RankingRow(keyword="premium body care products", position=Evidence.verified("3", "Sheet"), change=Evidence.verified("1", "Sheet"), competition="medium"),
            RankingRow(keyword="organic essential oils", position=Evidence.verified("5", "Sheet"), change=Evidence.verified("-1", "Sheet"), competition="medium"),
            RankingRow(keyword="natural fragrance for men", position=Evidence.verified("1", "Sheet"), change=Evidence.verified("2", "Sheet"), competition="low"),
            RankingRow(keyword="best attar perfume online", position=Evidence.verified("8", "Sheet"), change=Evidence.verified("0", "Sheet"), competition="high"),
            RankingRow(keyword="sandalwood oil for skin", position=Evidence.verified("12", "Sheet"), change=Evidence.verified("3", "Sheet"), competition="medium"),
            RankingRow(keyword="rose water face mist", position=Evidence.verified("4", "Sheet"), change=Evidence.verified("-2", "Sheet"), competition="low"),
            RankingRow(keyword="herbal shampoo natural", position=Evidence.verified("7", "Sheet"), change=Evidence.verified("1", "Sheet"), competition="medium"),
        ],
        kpis=Kpidata(
            clicks=Evidence.verified("2450", "Sheet:Traffic"),
            impressions=Evidence.verified("45200", "Sheet:Traffic"),
            clicks_change=Evidence.verified("12", "Sheet:Traffic"),
            impressions_change=Evidence.verified("8", "Sheet:Traffic"),
        ),
        action_plan=[
            ActionItem(team="Dev", task="Fix broken 404 links on /about page", priority="P0", impact="high", effort="2h", eta="Week 1", owner="John", status="todo"),
            ActionItem(team="SEO", task="Add missing H1 tags to 2 pages", priority="P1", impact="high", effort="1h", eta="Week 1", owner="Jane", status="todo"),
            ActionItem(team="Content", task="Rewrite thin content on /contact (target 500+ words)", priority="P1", impact="medium", effort="3h", eta="Week 2", owner="Mike", status="todo"),
            ActionItem(team="SEO", task="Build 5 new backlinks from industry directories", priority="P2", impact="medium", effort="8h", eta="Week 3", owner="Jane", status="todo"),
        ],
        executive_narrative="Test Client Inc. showed solid organic growth this month with 2 new #1 rankings and 12% growth in clicks. Mobile Core Web Vitals need attention (score: 62/100). Backlink profile: 156 links from 42 domains.",
    )
    f.rankings_total_estimated = 25
    return f

def get_all_text(docx_path):
    doc = Document(docx_path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)

def check_content(text, checks, label):
    missing = [c for c in checks if c.lower() not in text.lower()]
    if missing:
        print(f"  [MISS] {label}: missing {missing}")
        return False
    print(f"  [OK]   {label}: all {len(checks)} content checks passed")
    return True

def test_all():
    facts = make_facts()

    # Generate
    paths = {
        "Technical Audit": os.path.join(OUT, "Test_Client_Technical_Audit.docx"),
        "Monthly Report": os.path.join(OUT, "Test_Client_Monthly_Report.docx"),
        "Action Plan": os.path.join(OUT, "Test_Client_Action_Plan.docx"),
    }
    builders = {
        "Technical Audit": lambda p: build_technical_audit(facts, p),
        "Monthly Report": lambda p: build_clients_report(facts, p),
        "Action Plan": lambda p: build_action_plan_docx(facts, p),
    }
    sizes = {}
    for label, fn in builders.items():
        fn(paths[label])
        sizes[label] = os.path.getsize(paths[label])
        print(f"  [GEN]  {label:25s}  {os.path.basename(paths[label]):45s}  {sizes[label]/1024:8.1f} KB")

    # Verify content
    print()
    all_ok = True

    # Technical Audit checks
    t = get_all_text(paths["Technical Audit"])
    all_ok &= check_content(t, [
        "TECHNICAL SEO AUDIT", "Score", "Core SEO Metrics", "Title",
        "Meta Description", "Core Web Vitals", "Desktop", "Mobile",
        "Backlink Profile", "Total Backlinks", "Issues Log",
        "Missing H1", "Broken link", "Recommendations",
    ], "Technical Audit")

    # Monthly Report checks
    r = get_all_text(paths["Monthly Report"])
    all_ok &= check_content(r, [
        "MONTHLY SEO REPORT", "Executive Dashboard", "Core SEO Metrics",
        "Backlinks", "Authority", "Total Backlinks", "Referring Domains",
        "Core Web Vitals", "Mobile", "Desktop",
        "Search Visibility", "Keywords", "Top Performing",
        "Technical SEO Health", "Recommendations",
        "Action Plan", "Missing H1", "Broken link",
    ], "Monthly Report")

    # Action Plan checks
    a = get_all_text(paths["Action Plan"])
    all_ok &= check_content(a, [
        "MONTHLY ACTION PLAN", "Action Plan", "P0", "P1", "P2",
        "Dev", "SEO", "Content", "Broken", "Missing H1", "backlinks",
    ], "Action Plan")

    print()
    print("=" * 60)
    verdict = "PASS" if all_ok else "FAIL - content checks missing"
    print(f"  VERDICT: {verdict}")
    print("=" * 60)
    return all_ok

if __name__ == "__main__":
    ok = test_all()
    sys.exit(0 if ok else 1)
