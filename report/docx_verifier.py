"""Post-generation DOCX quality verifier.

Opens each generated DOCX and checks for empty/placeholder content so
inconsistencies and missing data are caught before the user downloads.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

PLACEHOLDER_PATTERNS: list[re.Pattern] = [
    re.compile(r"^\u2014$"),           # em-dash
    re.compile(r"^—$"),                # em-dash alt
    re.compile(r"^N/A$", re.IGNORECASE),
    re.compile(r"^Not Found$", re.IGNORECASE),
    re.compile(r"^unavailable$", re.IGNORECASE),
    re.compile(r"^None$", re.IGNORECASE),
    re.compile(r"^\s*$"),              # whitespace only
    re.compile(r"^0$"),                # zero (suspicious for some fields)
]

MAX_SUSPICIOUS_ZERO_LENGTH = 3  # only flag "0" if text is exactly "0"


def _is_empty_or_placeholder(text: str) -> bool:
    t = text.strip()
    if not t:
        return True
    return any(p.match(t) for p in PLACEHOLDER_PATTERNS)


def _text_from_paragraph(p) -> str:
    return "".join(run.text for run in p.runs if run.text)


def verify_docx(filepath: str | Path) -> dict[str, Any]:
    """Analyse a single DOCX file.

    Returns:
        {path, file_size_kb, total_paragraphs, empty_paragraphs,
         total_table_cells, empty_cells, sections_found, issues}
    """
    path = Path(filepath)
    if not path.exists():
        return {"path": str(path), "error": "File not found"}

    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        return {"path": str(path), "error": str(e)}

    total_paras = 0
    empty_paras = 0
    total_cells = 0
    empty_cells = 0
    sections: list[str] = []

    # Scan body paragraphs (non-empty headings → sections)
    for para in doc.paragraphs:
        text = _text_from_paragraph(para)
        total_paras += 1
        if _is_empty_or_placeholder(text):
            empty_paras += 1
        elif para.style.name.startswith("Heading"):
            sections.append(text[:80])

    # Scan table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_cells += 1
                cell_text = cell.text.strip()
                if _is_empty_or_placeholder(cell_text):
                    empty_cells += 1

    # Compute quality score (0-100)
    total_items = total_paras + total_cells
    empty_items = empty_paras + empty_cells
    quality = 100
    if total_items > 0:
        empty_ratio = empty_items / total_items
        quality = max(0, 100 - int(empty_ratio * 100))

    issues: list[str] = []
    if empty_paras > total_paras * 0.3:
        issues.append(f"{empty_paras}/{total_paras} paragraphs are empty/placeholder")
    if empty_cells > total_cells * 0.3:
        issues.append(f"{empty_cells}/{total_cells} table cells are empty/placeholder")
    if not sections:
        issues.append("No heading sections detected")

    result = {
        "path": str(path),
        "file_size_kb": round(path.stat().st_size / 1024, 1),
        "total_paragraphs": total_paras,
        "empty_paragraphs": empty_paras,
        "total_table_cells": total_cells,
        "empty_cells": empty_cells,
        "quality_score": quality,
        "sections_found": sections,
        "issues": issues,
        "passed": quality >= 50 and len(issues) == 0,
    }
    logger.info("Verifier: %s — quality=%d/100, sections=%d, issues=%s",
                 path.name, quality, len(sections), issues or "none")
    return result


def verify_all(filepaths: list[str | Path]) -> list[dict[str, Any]]:
    """Run verification on multiple DOCX files and return combined results."""
    return [verify_docx(fp) for fp in filepaths]
