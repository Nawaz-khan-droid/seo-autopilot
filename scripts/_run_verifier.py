"""Run the verifier on all existing reports and print results + content inspection."""

from __future__ import annotations

import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from pathlib import Path
from docx import Document as DocxDocument
from report.docx_verifier import verify_all

out = Path("output")

files = [
    out / "Digichefs Technical SEO Audit June 2026.docx",
    out / "Digichefs Monthly SEO Report June 2026.docx",
    out / "Digichefs SEO Action Plan June 2026.docx",
]

existing = [f for f in files if f.exists()]
print(f"Found {len(existing)}/{len(files)} Digichefs reports\n")

# 1. Verifier quality check
for r in verify_all(existing):
    print(f"=== {r['path']} ===")
    print(f"  Quality Score : {r['quality_score']}/100")
    print(f"  Passed        : {r['passed']}")
    print(f"  File Size     : {r['file_size_kb']} KB")
    print(f"  Paragraphs    : {r['total_paragraphs']} total, {r['empty_paragraphs']} empty")
    print(f"  Table Cells   : {r['total_table_cells']} total, {r['empty_cells']} empty")
    print(f"  Sections      : {r['sections_found']}")
    if r["issues"]:
        for i in r["issues"]:
            print(f"  [ISSUE] {i}")
    print()

# 2. Deep content inspection
print("=" * 60)
print("DEEP CONTENT INSPECTION -- real vs estimated metrics")
print("=" * 60)

for fp in existing:
    print(f"\n--- {fp.name} ---")
    doc = DocxDocument(str(fp))

    # Print all paragraph content
    print("  [PARAGRAPHS]")
    for p in doc.paragraphs:
        txt = "".join(r.text for r in p.runs if r.text).strip()
        if txt:
            print(f"    {txt[:130]}")

    # Print table data
    tables = doc.tables
    print(f"  [TABLES: {len(tables)} table(s)]")
    for ti, table in enumerate(tables):
        print(f"    Table {ti+1}: {len(table.rows)} rows x {len(table.columns)} cols")
        for row in table.rows:
            cells = [cell.text.strip()[:30] for cell in row.cells]
            print(f"      | {'  |  '.join(cells)} |")

print("\nDone.")
