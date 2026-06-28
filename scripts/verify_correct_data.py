"""Verify the report uses Beautiful India's correct data, not DigiChefs."""
import os
from docx import Document

d = r"C:\Users\nawaz\OneDrive\Desktop\SEO Demo\seo-autopilot\output"
docx_files = [f for f in os.listdir(d) if f.startswith("Beautiful_India") and f.endswith(".docx") and "v2" not in f]
latest = sorted(docx_files)[-1]
fp = os.path.join(d, latest)

doc = Document(fp)

print(f"=== VERIFYING: {latest} ({os.path.getsize(fp)} bytes) ===")

# Get ALL text including tables
all_text = ""
for p in doc.paragraphs:
    all_text += p.text + "\n"
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            all_text += cell.text + " "

# Check for WRONG keywords (DigiChefs)
wrong_keywords = [
    "search engine optimization company mumbai",
    "seo services mumbai",
    "best seo company mumbai",
    "seo agency mumbai",
    "digital marketing agency mumbai",
    "content marketing strategy",
    "website designing company mumbai",
]
print("\nCHECK: DigiChefs keywords ABSENT")
all_clean = True
for kw in wrong_keywords:
    if kw in all_text.lower():
        print(f"  FAIL: '{kw}' found!")
        all_clean = False
if all_clean:
    print("  PASS: No DigiChefs keywords")

# Check for CORRECT keywords (Beautiful India)
correct_keywords = [
    "beautiful india perfume",
    "luxury perfume india",
    "beautiful india one",
    "discovery set",
    "luxury body care",
    "peace perfume",
]
print("\nCHECK: Beautiful India keywords PRESENT")
found_any = False
for kw in correct_keywords:
    if kw.lower() in all_text.lower():
        print(f"  OK: '{kw}' found")
        found_any = True
if not found_any:
    print("  WARNING: None of expected Beautiful India keywords found")

# Check for fake names
print("\nCHECK: No fake names")
for name in ["Raj", "Priya", "Anika", "Vikram"]:
    if name in all_text:
        print(f"  FAIL: '{name}' found!")
        all_clean = False
if all_clean:
    print("  PASS: No fake names")

# Check real products mentioned
product_checks = ["One", "Discovery Set", "perfume", "luxury"]
print("\nCHECK: Product context")
for p in product_checks:
    if p.lower() in all_text.lower():
        print(f"  OK: '{p}' referenced")

print(f"\n=== VERDICT: {'PASS' if all_clean else 'ISSUES FOUND'} ===")
