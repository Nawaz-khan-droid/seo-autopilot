"""Monthly Action Plan — Internal Team DOCX (6-10 slide cards).

Smart Advisor/Analyst/Supervisor → Internal Team tone.
Contains technical details (CWV metrics, technical specs) that are
deliberately excluded from the client-facing report.
"""

from __future__ import annotations

import os
import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from report.facts import ReportFacts

logger = logging.getLogger(__name__)

# ── Palette ──
NAVY = RGBColor(0x0F, 0x27, 0x47)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
DARK = RGBColor(0x1E, 0x29, 0x3B)
BODY = RGBColor(0x33, 0x41, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
RED = RGBColor(0xDC, 0x26, 0x26)
TEAL = RGBColor(0x14, 0xB8, 0xA6)
GRAY = RGBColor(0x94, 0xA3, 0xB8)

FN = "Calibri"

def _shd(cell, c):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _cell(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FN
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align is not None: p.alignment = align

def _para(doc, text, bold=False, color=None, size=10, align=None, before=0, after=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FN
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p

def _ev(v):
    if v is None: return "\u2014"
    if hasattr(v, "is_available") and not v.is_available: return "\u2014"
    val = v.value if hasattr(v, "value") else v
    return str(val) if val is not None else "\u2014"

def _evi(v):
    try: return int(float(_ev(v)))
    except: return 0

def _sev_color(s):
    s = s.lower().strip()
    if s == "critical": return RED
    if s in ("high", "error"): return AMBER
    if s in ("medium", "warning"): return ACCENT
    return GRAY

# ═══════════════════════════════════════════════════════════════

def build_action_plan_docx(facts: ReportFacts, output_path: str = "") -> str:
    """Build internal team action plan DOCX (6-10 slide cards)."""
    m = facts.metadata
    if not output_path:
        sn = (m.client_name or "Client").replace(" ", "_")
        sm = (m.report_month or "Report").replace(" ", "_")
        output_path = f"{sn}_{sm}_Action_Plan.docx"

    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = FN
    style.font.size = Pt(10)
    style.font.color.rgb = BODY
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.2

    for lvl, (sz, clr, sb) in {1: (17, NAVY, 4), 2: (13, ACCENT, 3)}.items():
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = FN
        hs.font.size = Pt(sz)
        hs.font.color.rgb = clr
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(sb)
        hs.paragraph_format.space_after = Pt(2)

    ranks = facts.rankings or []
    total = len(ranks)
    improved = sum(1 for r in ranks if _evi(r.change) > 0)
    dropped = sum(1 for r in ranks if _evi(r.change) < 0)

    # ════════════════════════════════════════════
    # CARD 1: COVER
    # ════════════════════════════════════════════
    for _ in range(6): doc.add_paragraph()
    _para(doc, "MONTHLY ACTION PLAN", bold=True, color=NAVY, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "\u2500" * 30, color=ACCENT, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "", size=4)
    if m.client_name:
        _para(doc, m.client_name, bold=True, color=NAVY, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    if m.report_month:
        _para(doc, m.report_month, color=BODY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "", size=8)
    _para(doc, "INTERNAL TEAM — CONFIDENTIAL", bold=True, color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"Prepared by: {m.agency_name or 'SEO Team'}", color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    gen = m.generated_at or datetime.now().strftime("%B %d, %Y")
    _para(doc, f"Generated: {gen}", color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # STRATEGIC OVERVIEW (LLM-generated, optional)
    # ════════════════════════════════════════════
    if facts.executive_narrative:
        _para(doc, "Strategic Overview", bold=True, color=NAVY, size=13)
        _para(doc, facts.executive_narrative, size=10, color=BODY)
        _para(doc, "", size=4)

    # ════════════════════════════════════════════
    # CARD 2: TASK LANDSCAPE
    # ════════════════════════════════════════════
    doc.add_heading("Task Landscape", level=1)

    _para(doc, f"Action items derived from this period's SEO analysis. "
           f"{improved} keywords improved, {dropped} dropped out of {total} tracked.", size=10)

    teams = ["SEO", "Dev", "Content", "Local"]
    team_items = {t: [a for a in facts.action_plan if a.team == t] for t in teams} if facts.action_plan else {}

    if team_items:
        lt = doc.add_table(rows=1 + len(teams), cols=4)
        lt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Team", "Task Count", "Top Priority", "Key ETA"]):
            _cell(lt.rows[0].cells[ci], h, bold=True, color=WHITE, size=9)
            _shd(lt.rows[0].cells[ci], "0F2747")
        for ri, team in enumerate(teams):
            items = team_items.get(team, [])
            top_p = min(items, key=lambda a: a.priority).priority if items else "\u2014"
            earliest = min(items, key=lambda a: a.eta).eta if items else "\u2014"
            _cell(lt.rows[ri+1].cells[0], team, bold=True, size=9)
            _cell(lt.rows[ri+1].cells[1], str(len(items)), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(lt.rows[ri+1].cells[2], top_p, size=9, bold=True,
                  color=RED if top_p in ("P0", "P1") else AMBER, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(lt.rows[ri+1].cells[3], earliest, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2 == 0:
                for ci in range(4): _shd(lt.rows[ri+1].cells[ci], "F8FAFC")

    # ════════════════════════════════════════════
    # CARD 3-6: TEAM TABLES (detailed)
    # ════════════════════════════════════════════
    team_labels = {
        "SEO": ("SEO Team Tasks", "Keyword optimization, technical fixes, monitoring"),
        "Dev": ("Development Team Tasks", "CWV optimization, site performance, technical implementation"),
        "Content": ("Content Team Tasks", "Content creation, optimization, gap analysis"),
        "Local": ("Local SEO Tasks", "Local listings, reviews, geo-targeted content"),
    }

    for team in teams:
        items = team_items.get(team, [])
        if not items:
            continue

        doc.add_page_break()
        lbl, desc = team_labels.get(team, (f"{team} Tasks", ""))
        doc.add_heading(lbl, level=1)
        if desc:
            _para(doc, desc, size=9, color=GRAY)

        # Advisor note — built from actual facts data, not hardcoded
        c = facts.cwv
        t = facts.technical
        si = facts.site_info
        cm = _evi(c.mobile_score)
        cd = _evi(c.desktop_score)
        lcp = _ev(c.lcp_seconds)

        note = ""
        pa = _evi(t.pages_audited)
        if team == "Dev":
            parts = []
            if cm < 90 or cd < 90:
                parts.append(f"Mobile PSI is {cm}/100, Desktop is {cd}/100")
            if lcp != "\u2014":
                try:
                    lcp_sec = float(lcp.rstrip("s"))
                    if lcp_sec > 2.5:
                        parts.append(f"LCP at {lcp_sec}s exceeds the 2.5s target — optimise hero images and server response")
                    else:
                        parts.append(f"LCP at {lcp_sec}s is within target")
                except Exception:
                    pass
            if _ev(c.cls_score) != "\u2014":
                parts.append(f"CLS at {_ev(c.cls_score)} needs monitoring")
            if pa > 0:
                parts.append(f"{pa} pages crawled for technical analysis")
            if not parts:
                parts.append("No PageSpeed or CWV data collected — connect PSI API key or run Playwright to generate recommendations.")
            note = "; ".join(parts) + "."
        elif team == "Content":
            parts = []
            wc = _evi(si.word_count)
            if wc > 0 and wc < 300:
                parts.append(f"Homepage is only {wc} words — expand to 500+ for topical relevance")
            elif wc > 0:
                parts.append(f"Page word count ({wc}) is adequate")
            mm = _evi(t.missing_meta)
            if mm > 0:
                parts.append(f"{mm} pages are missing meta descriptions across {pa} pages crawled")
            mh = _evi(t.missing_h1)
            if mh > 0:
                parts.append(f"{mh} pages have missing or duplicate H1 tags across {pa} pages")
            if not parts:
                parts.append("No crawl data available — content recommendations require at least 1 page to be analyzed.")
            note = "; ".join(parts) + "."
        elif team == "SEO":
            parts = []
            total_iss = _evi(t.total_issues)
            if total_iss > 0:
                parts.append(f"{total_iss} technical issues found across {pa} pages — prioritise by severity")
            if si.has_robots_txt.is_available and "Not Found" in str(si.has_robots_txt.value):
                parts.append("robots.txt is missing — search engines may waste crawl budget")
            if si.has_sitemap_xml.is_available and "Not Found" in str(si.has_sitemap_xml.value):
                parts.append("sitemap.xml is missing — new pages may not be discovered")
            if not si.has_og_tags.is_available or "No" in str(si.has_og_tags.value):
                parts.append("Open Graph tags are missing — social shares will not render rich previews")
            has_sch = t.has_schema.is_available and "Yes" in str(t.has_schema.value)
            if not has_sch:
                parts.append("schema.org markup is absent — add structured data for rich results eligibility")
            if not parts:
                if pa == 0:
                    parts.append("No crawl data collected — SEO recommendations require at least 1 analyzed page.")
                else:
                    parts.append("robots.txt, sitemap.xml, OG tags, and schema were all detected as expected.")
            note = "; ".join(parts) + "."
        elif team == "Local":
            has_local_data = (facts.local_seo.map_pack_presence.is_available or
                              facts.local_seo.review_count.is_available)
            if has_local_data:
                parts = []
                if _ev(facts.local_seo.review_count) != "\u2014":
                    parts.append(f"GMB reviews: {_ev(facts.local_seo.review_count)}")
                if _ev(facts.local_seo.avg_rating) != "\u2014":
                    parts.append(f"Avg rating: {_ev(facts.local_seo.avg_rating)}")
                note = "; ".join(parts) + "." if parts else "Local data available — see report for details."
            else:
                note = "No local SEO data was provided this cycle. To enable Local actions, configure GBP tracking or upload data."

        if note:
            _para(doc, f"Advisor Note: {note}", size=9, color=ACCENT, before=4)

        # Table: Task | Priority | ETA
        tt = doc.add_table(rows=1 + len(items), cols=3)
        tt.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Task", "Priority", "ETA"]):
            _cell(tt.rows[0].cells[ci], h, bold=True, color=WHITE, size=8)
            _shd(tt.rows[0].cells[ci], "0F2747")
        for ri, a in enumerate(items):
            pc = RED if a.priority in ("P0", "P1") else (AMBER if a.priority == "P2" else GRAY)
            _cell(tt.rows[ri+1].cells[0], a.task[:80], size=8)
            _cell(tt.rows[ri+1].cells[1], a.priority, size=8, bold=True, color=pc, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(tt.rows[ri+1].cells[2], a.eta, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2 == 0:
                for ci in range(3): _shd(tt.rows[ri+1].cells[ci], "F8FAFC")

    # ════════════════════════════════════════════
    # CARD 7: PRIORITY MATRIX
    # ════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("Priority Matrix", level=1)

    _para(doc, "All tasks mapped by priority level for sprint planning:", size=10)

    for p_lbl, p_filter in [("P0 — Critical (This Week)", lambda a: a.priority == "P0"),
                              ("P1 — High (Next Week)", lambda a: a.priority == "P1"),
                              ("P2 — Medium (This Sprint)", lambda a: a.priority == "P2"),
                              ("P3 — Low (Backlog)", lambda a: a.priority == "P3")]:
        items = [a for a in facts.action_plan if p_filter(a)] if facts.action_plan else []
        if items:
            doc.add_heading(p_lbl, level=2)
            for a in items[:5]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(f"[{a.team}] {a.task[:80]}")
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = FN
                p.add_run(f"  ({a.effort}, {a.eta})").font.size = Pt(8)

    # ════════════════════════════════════════════
    # CARD 8: TECHNICAL SPECIFICATIONS (advisor detail)
    # ════════════════════════════════════════════
    c = facts.cwv
    cm = _evi(c.mobile_score)
    cd = _evi(c.desktop_score)
    lcp = _ev(c.lcp_seconds)
    inp = _ev(c.inp_ms)
    cls_val = _ev(c.cls_score)
    rbm = _ev(c.render_blocking_ms)

    if any(x != "\u2014" for x in [lcp, inp, cls_val, rbm]):
        doc.add_page_break()
        doc.add_heading("Technical Specifications — Dev Team", level=1)
        _para(doc, "Detailed Core Web Vitals and performance metrics for the development team. "
               "These are excluded from the client report.", size=9, color=GRAY)

        spec_data = [
            ("Mobile PSI Score", f"{cm}/100", "Target >90"),
            ("Desktop PSI Score", f"{cd}/100", "Target >90"),
            ("LCP (Largest Contentful Paint)", f"{lcp}s", "Target <2.5s"),
            ("INP (Interaction to Next Paint)", f"{inp}ms", "Target <200ms"),
            ("CLS (Cumulative Layout Shift)", cls_val, "Target <0.1"),
            ("Render-blocking Resources", f"{rbm}ms", "Target <500ms"),
        ]
        st = doc.add_table(rows=1 + len(spec_data), cols=3)
        st.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(["Metric", "Current", "Target"]):
            _cell(st.rows[0].cells[ci], h, bold=True, color=WHITE, size=8)
            _shd(st.rows[0].cells[ci], "0F2747")
        for ri, (lb, cur, tgt) in enumerate(spec_data):
            _cell(st.rows[ri+1].cells[0], lb, size=8)
            _cell(st.rows[ri+1].cells[1], cur, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(st.rows[ri+1].cells[2], tgt, size=8, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2 == 0:
                for ci in range(3): _shd(st.rows[ri+1].cells[ci], "F8FAFC")

        _para(doc, "", size=4)
        priority_parts = []
        if lcp != "\u2014":
            try:
                lcp_sec = float(lcp.rstrip("s"))
                if lcp_sec > 2.5:
                    priority_parts.append(f"LCP at {lcp_sec}s exceeds 2.5s target")
            except:
                pass
        if cls_val != "\u2014":
            priority_parts.append(f"CLS at {cls_val}")
        if rbm != "\u2014":
            priority_parts.append(f"render-blocking resources at {rbm}ms")
        if priority_parts:
            _para(doc, f"Priority: {'; '.join(priority_parts)}.", size=9, color=BODY)

    # ════════════════════════════════════════════
    # CARD 9: ADVISOR NOTES & NEXT STEPS
    # ════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("Advisor Summary", level=1)

    _para(doc, "Key takeaways for the team lead:", bold=True, size=11)
    _para(doc, "", size=4)

    takeaways = []
    if total > 0:
        if improved > 0:
            takeaways.append(f"Positive momentum: {improved} keywords improved this period — maintain current strategy")
        if dropped > 0:
            takeaways.append(f"Watch list: {dropped} keywords dropped — assign content team for recovery")
    if cm > 0 and cm < 80:
        takeaways.append(f"Mobile experience is the top technical priority (score {cm}/100) — assign Dev team to CWV sprint")
    if facts.technical.issues_list and any(iss.severity.lower() == "critical" for iss in facts.technical.issues_list):
        takeaways.append("Critical technical issues need immediate attention — SEO + Dev collaboration required")
    if not takeaways:
        takeaways.append("No actionable data was collected this cycle — confirm crawl and API configurations before next sprint.")

    for tw in takeaways:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(tw)
        r.font.size = Pt(10)
        r.font.name = FN

    _para(doc, "", size=6)
    if total > 0 or cm > 0 or facts.technical.issues_list:
        _para(doc, "Next sprint planning session: Review priorities, assign owners, confirm ETAs.", size=10, bold=True)

    # ── Footer ──
    for sec in doc.sections:
        ftr = sec.footer
        ftr.is_linked_to_previous = False
        p = ftr.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("\u2500" * 50)
        r.font.size = Pt(5)
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
        p2 = ftr.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(f"INTERNAL ONLY — {m.agency_name or 'SEO Team'}  |  {m.client_name or 'Client'}  |  {m.report_month or ''}")
        r.font.size = Pt(7)
        r.font.color.rgb = GRAY
        r.font.name = FN

    doc.save(output_path)
    return output_path
